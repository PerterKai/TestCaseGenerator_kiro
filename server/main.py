#!/usr/bin/env python3
"""
Test Case Generator MCP Server v7.0
- Phase-based workflow with file cache for cross-session resume
- Splits document reading into summary + sections to reduce context
- External multimodal LLM API support with GUI configuration
- Multi-threaded image processing
- Improved image extraction (VML, OLE, orphan detection)
- All state persisted to .tmp/cache/ for recovery
"""

import sys
import json
import os
import glob
import zipfile
import re
import subprocess
import base64
import hashlib
import shutil
import threading
import traceback
from io import BytesIO
from xml.etree import ElementTree as ET

# ============================================================
# Windows binary mode for stdin/stdout
# ============================================================
if sys.platform == "win32":
    import msvcrt
    msvcrt.setmode(sys.stdin.fileno(), os.O_BINARY)
    msvcrt.setmode(sys.stdout.fileno(), os.O_BINARY)

sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# ============================================================
# Auto-install helper
# ============================================================

def _ensure_pkg(import_name, pip_name):
    try:
        __import__(import_name)
        return True
    except ImportError:
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", pip_name, "-q"],
                stdout=subprocess.DEVNULL, stderr=subprocess.PIPE)
            return True
        except Exception:
            return False

# ============================================================
# MCP Protocol (stdio JSON-RPC)
# ============================================================

def _send_msg(msg_bytes):
    """Write a JSON-RPC message with Content-Length header (MCP stdio transport spec)."""
    header = f"Content-Length: {len(msg_bytes)}\r\n\r\n".encode('ascii')
    sys.stdout.buffer.write(header + msg_bytes)
    sys.stdout.buffer.flush()


def send_response(rid, result):
    msg = json.dumps({"jsonrpc": "2.0", "id": rid, "result": result}, ensure_ascii=False)
    _send_msg(msg.encode('utf-8'))


def send_error(rid, code, message):
    msg = json.dumps({"jsonrpc": "2.0", "id": rid, "error": {"code": code, "message": message}}, ensure_ascii=False)
    _send_msg(msg.encode('utf-8'))


# ============================================================
# Constants & Paths
# ============================================================

# Resolve workspace root: Kiro sets cwd to workspace root when launching MCP.
# Also support --workspace CLI arg for manual runs.
def _resolve_initial_workspace():
    """Determine workspace root directory at startup."""
    # Check CLI args first
    for i, arg in enumerate(sys.argv):
        if arg == "--workspace" and i + 1 < len(sys.argv):
            return os.path.abspath(sys.argv[i + 1])
    # Default: cwd (Kiro sets this to workspace root)
    return os.getcwd()

_INITIAL_WORKSPACE = _resolve_initial_workspace()

WORKSPACE_DIR = None
TMP_DOC_DIR = os.path.join(_INITIAL_WORKSPACE, ".tmp", "doc_mk")
TMP_PIC_DIR = os.path.join(_INITIAL_WORKSPACE, ".tmp", "picture")
TMP_CACHE_DIR = os.path.join(_INITIAL_WORKSPACE, ".tmp", "cache")

# No forced session switch — let the system decide naturally.
# Cross-session resume is still fully supported via .tmp/cache/.


def _update_workspace(directory):
    global WORKSPACE_DIR, TMP_DOC_DIR, TMP_PIC_DIR, TMP_CACHE_DIR
    WORKSPACE_DIR = directory
    TMP_DOC_DIR = os.path.join(directory, ".tmp", "doc_mk")
    TMP_PIC_DIR = os.path.join(directory, ".tmp", "picture")
    TMP_CACHE_DIR = os.path.join(directory, ".tmp", "cache")


def _workspace():
    return WORKSPACE_DIR or _INITIAL_WORKSPACE

# ============================================================
# Cache filename constants
# ============================================================
CACHE_PHASE_STATE = "phase_state.json"
CACHE_IMAGE_PROGRESS = "image_progress.json"
CACHE_TESTCASES = "testcases.json"
CACHE_DOC_SUMMARY = "doc_summary.json"

# ============================================================
# Cache / Persistence Layer
# ============================================================

def _cache_path(filename):
    return os.path.join(TMP_CACHE_DIR, filename)


def _save_cache(filename, data):
    os.makedirs(TMP_CACHE_DIR, exist_ok=True)
    path = _cache_path(filename)
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        sys.stderr.write(f"[MCP] Warning: failed to save cache {filename}: {e}\n")
        sys.stderr.flush()


def _load_cache(filename, default=None):
    path = _cache_path(filename)
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            sys.stderr.write(f"[MCP] Warning: failed to load cache {filename}: {e}\n")
            sys.stderr.flush()
    return default if default is not None else {}


def _save_phase_state(phase, status, extra=None):
    """Update phase_state.json with current progress."""
    state = _load_cache(CACHE_PHASE_STATE, {
        "current_phase": phase,
        "workspace_dir": _workspace(),
        "phases": {}
    })
    state["current_phase"] = phase
    state["workspace_dir"] = _workspace()
    state.setdefault("phases", {})
    if phase not in state["phases"]:
        state["phases"][phase] = {}
    state["phases"][phase]["status"] = status
    if extra:
        state["phases"][phase].update(extra)
    _save_cache(CACHE_PHASE_STATE, state)
    return state


def _reset_phase_state():
    """Clear all phase state for a fresh start."""
    state = {
        "current_phase": "init",
        "workspace_dir": _workspace(),
        "phases": {}
    }
    _save_cache(CACHE_PHASE_STATE, state)
    return state

# ============================================================
# Document tagging constants
# ============================================================
# Tags in filename that mark a document as "primary" (用例生成目标)
# Primary docs get full processing: text + images + test case generation
# Other docs are "reference" (辅助资料): text only, consulted on demand
PRIMARY_DOC_TAGS = ['【主prd】', '【主概设】', '【主后端概设】', '【主前端概设】']

def _classify_document(filename):
    """Classify a document as 'primary' or 'reference' based on filename tags.
    
    Returns ('primary', tag) if filename contains a primary tag, else ('reference', None).
    If NO documents have primary tags, all are treated as primary (backward compat).
    """
    name_lower = filename.lower()
    for tag in PRIMARY_DOC_TAGS:
        if tag.lower() in name_lower:
            return 'primary', tag
    return 'reference', None

# ============================================================
# Image filtering constants
# ============================================================
# Only process images where: min(w,h) > 64 AND max(w,h) >= 224
IMG_MIN_SHORT_EDGE = 64   # shortest side must be > this
IMG_MIN_LONG_EDGE = 224   # longest side must be >= this

def _should_process_image(w, h):
    """Return True if image dimensions qualify for analysis."""
    short = min(w, h)
    long_ = max(w, h)
    return short > IMG_MIN_SHORT_EDGE and long_ >= IMG_MIN_LONG_EDGE


# ============================================================
# Image helpers
# ============================================================

def _img_mime(ext):
    return {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
            "emf": "image/emf", "wmf": "image/wmf"}.get(ext.lower().lstrip('.'), "image/png")


def _generate_image_id(doc_name, img_name):
    raw = f"{doc_name}_{img_name}"
    short_hash = hashlib.md5(raw.encode()).hexdigest()[:8]
    base = os.path.splitext(img_name)[0]
    ext = os.path.splitext(img_name)[1]
    return f"{base}_{short_hash}{ext}"


def _resize_image(img_data, ext):
    final_data = img_data
    final_mime = _img_mime(ext)
    try:
        from PIL import Image
        img_obj = Image.open(BytesIO(img_data))
        w, h = img_obj.size

        # Skip images that don't meet dimension threshold
        if not _should_process_image(w, h):
            return final_data, final_mime

        # Resize large images, scale based on content density
        MAX_DIM = 1568  # Claude vision optimal tile boundary (multiple of 784)
        if w > MAX_DIM or h > MAX_DIM:
            ratio = min(MAX_DIM / w, MAX_DIM / h)
            new_w, new_h = int(w * ratio), int(h * ratio)
            img_obj = img_obj.resize((new_w, new_h), Image.LANCZOS)

        # Convert to grayscale for document images
        if img_obj.mode not in ('L', 'LA'):
            img_obj = img_obj.convert('L')

        buf = BytesIO()
        img_obj.save(buf, format='PNG', optimize=True)
        final_mime = "image/png"
        final_data = buf.getvalue()
    except Exception:
        pass
    return final_data, final_mime


def _resize_image_for_llm(img_data, mime):
    """Resize image for LLM vision API: max longest edge 3840, PNG grayscale."""
    try:
        from PIL import Image
        img_obj = Image.open(BytesIO(img_data))
        w, h = img_obj.size

        MAX_DIM = 3840
        if w > MAX_DIM or h > MAX_DIM:
            ratio = min(MAX_DIM / w, MAX_DIM / h)
            new_w, new_h = int(w * ratio), int(h * ratio)
            img_obj = img_obj.resize((new_w, new_h), Image.LANCZOS)

        # Convert to grayscale
        if img_obj.mode not in ('L', 'LA'):
            img_obj = img_obj.convert('L')

        buf = BytesIO()
        img_obj.save(buf, format='PNG', optimize=True)
        return buf.getvalue(), "image/png"
    except Exception:
        return img_data, mime

# ============================================================
# DOCX → Markdown + Images extraction
# ============================================================

def _build_rid_to_media(filepath):
    rid_to_media = {}
    header_footer_media = set()  # images only referenced by headers/footers (e.g. watermarks)
    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            # Collect all actual media files in the zip for cross-reference
            actual_media = set()
            for name in z.namelist():
                if name.startswith('word/media/'):
                    actual_media.add(os.path.basename(name))

            # Track which media are referenced by document.xml vs headers/footers
            doc_rels_media = set()
            hf_rels_media = set()

            # Parse all .rels files under word/ (document, headers, footers, etc.)
            rels_files = [n for n in z.namelist()
                          if n.startswith('word/') and n.endswith('.rels')]
            for rels_path in rels_files:
                # Determine if this rels file belongs to a header/footer
                rels_basename = os.path.basename(rels_path).replace('.rels', '')
                is_hf = any(k in rels_basename for k in ('header', 'footer'))
                try:
                    rels_root = ET.fromstring(z.read(rels_path))
                    for rel in rels_root:
                        target = rel.get('Target', '')
                        rid = rel.get('Id', '')
                        rel_type = rel.get('Type', '')
                        media_name = None
                        # Match media/ in target path (handles media/xxx and ../media/xxx)
                        if 'media/' in target:
                            media_name = target.split('/')[-1]
                            rid_to_media[rid] = media_name
                        # Also match image relationship types without media/ in path
                        elif 'image' in rel_type.lower():
                            basename = target.split('/')[-1]
                            if basename in actual_media:
                                media_name = basename
                                rid_to_media[rid] = media_name
                        # OLE object images
                        elif 'oleObject' in rel_type or 'package' in rel_type.lower():
                            basename = target.split('/')[-1]
                            if basename in actual_media:
                                media_name = basename
                                rid_to_media[rid] = media_name

                        if media_name:
                            if is_hf:
                                hf_rels_media.add(media_name)
                            else:
                                doc_rels_media.add(media_name)
                except Exception:
                    continue

            # Images referenced ONLY by headers/footers (not by document body) are watermarks
            header_footer_media = hf_rels_media - doc_rels_media

            # Ensure all media files are discoverable: create reverse mapping
            # for any media file not yet referenced by a relationship
            referenced_media = set(rid_to_media.values())
            unreferenced = actual_media - referenced_media
            if unreferenced:
                sys.stderr.write(f"[MCP] Found {len(unreferenced)} unreferenced media files: {unreferenced}\n")
                sys.stderr.flush()
    except Exception as e:
        sys.stderr.write(f"[MCP] Warning: failed to build rid_to_media for {filepath}: {e}\n")
        sys.stderr.flush()
    return rid_to_media, header_footer_media


def _build_ole_excel_map(filepath):
    """Build mapping: preview_image_rId -> xlsx_zip_path for embedded Excel objects.
    
    In docx, embedded Excel tables appear as w:object containing:
    - v:shape > v:imagedata (preview image, rId -> media/imageX.png)
    - o:OLEObject with ProgID=Excel.Sheet.12 (rId -> embeddings/xxx.xlsx)
    
    Returns dict: {preview_image_media_name: xlsx_zip_path}
    """
    ole_map = {}  # media_name -> xlsx_zip_path
    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            if 'word/document.xml' not in z.namelist():
                return ole_map
            
            # Build rId -> target mapping from rels
            rid_to_target = {}
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path in z.namelist():
                rels_root = ET.fromstring(z.read(rels_path))
                for rel in rels_root:
                    rid_to_target[rel.get('Id', '')] = rel.get('Target', '')
            
            # Parse document.xml for w:object elements
            W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            
            doc_root = ET.fromstring(z.read('word/document.xml'))
            for obj in doc_root.iter(f'{{{W_NS}}}object'):
                ole_elem = None
                img_rid = None
                
                for child in obj:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'OLEObject':
                        prog_id = child.get('ProgID', '')
                        if prog_id.startswith('Excel.'):
                            ole_rid = child.get(f'{{{R_NS}}}id', '')
                            ole_elem = child
                    elif tag == 'shape':
                        for sub in child:
                            sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
                            if sub_tag == 'imagedata':
                                img_rid = sub.get(f'{{{R_NS}}}id', '')
                
                if ole_elem is not None and img_rid:
                    ole_rid = ole_elem.get(f'{{{R_NS}}}id', '')
                    ole_target = rid_to_target.get(ole_rid, '')
                    img_target = rid_to_target.get(img_rid, '')
                    
                    if ole_target and img_target:
                        # img_target is like "media/image1.png"
                        img_media_name = img_target.split('/')[-1]
                        # ole_target is like "embeddings/xxx.xlsx"
                        xlsx_path = 'word/' + ole_target if not ole_target.startswith('word/') else ole_target
                        if xlsx_path in z.namelist():
                            ole_map[img_media_name] = xlsx_path
                            sys.stderr.write(f"[MCP] OLE Excel detected: {img_media_name} -> {xlsx_path}\n")
                            sys.stderr.flush()
    except Exception as e:
        sys.stderr.write(f"[MCP] OLE Excel map error: {e}\n")
        sys.stderr.flush()
    return ole_map


def _parse_embedded_xlsx(z, xlsx_zip_path):
    """Parse an embedded xlsx file from the docx zip into markdown table(s).
    
    Returns markdown string with table content, or empty string on failure.
    """
    try:
        _ensure_pkg("openpyxl", "openpyxl")
        import openpyxl
        
        xlsx_data = z.read(xlsx_zip_path)
        wb = openpyxl.load_workbook(BytesIO(xlsx_data), data_only=True)
        md_parts = []
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            max_row = min(ws.max_row or 0, 200)  # Safety cap
            for row in ws.iter_rows(min_row=1, max_row=max_row, values_only=True):
                cells = [str(c).replace('\n', ' ').replace('|', '\\|') if c is not None else "" for c in row]
                # Trim trailing empty cells
                while cells and cells[-1] == "":
                    cells.pop()
                if cells:
                    rows.append(cells)
            
            if not rows:
                continue
            
            # Normalize column count
            max_cols = max(len(r) for r in rows)
            for r in rows:
                while len(r) < max_cols:
                    r.append("")
            
            # Build markdown table
            lines = []
            if len(wb.sheetnames) > 1:
                lines.append(f"**{sheet_name}**")
                lines.append("")
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row[:max_cols]) + " |")
            md_parts.append("\n".join(lines))
        
        wb.close()
        return "\n\n".join(md_parts)
    except Exception as e:
        sys.stderr.write(f"[MCP] xlsx parse error ({xlsx_zip_path}): {e}\n")
        sys.stderr.flush()
        return ""


def _find_images_in_element(elem, rid_to_media):
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    V_NS = 'urn:schemas-microsoft-com:vml'
    O_NS = 'urn:schemas-microsoft-com:office:office'
    found = []
    seen = set()

    def _add(rid):
        if rid and rid in rid_to_media:
            name = rid_to_media[rid]
            if name not in seen:
                seen.add(name)
                found.append(name)

    # 1. DrawingML blip (standard images)
    for blip in elem.iter(f'{{{A_NS}}}blip'):
        _add(blip.get(f'{{{R_NS}}}embed', ''))
        _add(blip.get(f'{{{R_NS}}}link', ''))

    # 2. VML w:pict > v:imagedata (legacy images)
    for pict in elem.iter(f'{{{W_NS}}}pict'):
        for child in pict.iter():
            _add(child.get(f'{{{R_NS}}}id', ''))
            _add(child.get(f'{{{R_NS}}}href', ''))
            _add(child.get(f'{{{R_NS}}}pict', ''))

    # 3. VML v:imagedata directly (some docs use VML namespace)
    for imgdata in elem.iter(f'{{{V_NS}}}imagedata'):
        _add(imgdata.get(f'{{{R_NS}}}id', ''))
        _add(imgdata.get(f'{{{R_NS}}}href', ''))
        for attr_name, attr_val in imgdata.attrib.items():
            if attr_name.endswith('}id') or attr_name.endswith('}href'):
                _add(attr_val)

    # 4. OLE objects with image representations
    for ole in elem.iter(f'{{{O_NS}}}OLEObject'):
        _add(ole.get(f'{{{R_NS}}}id', ''))

    # 5. Fallback: scan all elements for r:embed / r:id pointing to media
    for child in elem.iter():
        for attr_name, attr_val in child.attrib.items():
            if ('embed' in attr_name.lower() or attr_name.endswith('}id')) and attr_val in rid_to_media:
                name = rid_to_media[attr_val]
                if name not in seen:
                    seen.add(name)
                    found.append(name)

    return found



def _detect_heading_level(para):
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    if para.style and para.style.name and para.style.name.startswith('Heading'):
        lvl_str = para.style.name.replace('Heading ', '').replace('Heading', '1')
        try:
            return int(lvl_str)
        except ValueError:
            return 1
    try:
        pPr = para._element.find(f'{W_NS}pPr')
        if pPr is not None:
            outlineLvl = pPr.find(f'{W_NS}outlineLvl')
            if outlineLvl is not None:
                val = int(outlineLvl.get(f'{W_NS}val', '-1'))
                if val >= 0:
                    return val + 1
    except Exception:
        pass
    max_size = 0
    is_bold = False
    for run in para.runs:
        if run.bold:
            is_bold = True
        if run.font.size:
            sz = run.font.size / 12700
            if sz > max_size:
                max_size = sz
    if max_size >= 22:
        return 1
    if max_size >= 17 and is_bold:
        return 1
    if max_size >= 15 and is_bold:
        return 2
    if max_size >= 14 and is_bold:
        return 3
    return 0

def _table_to_markdown(table, rid_to_media, doc_name, image_registry):
    rows = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
            cell_images = []
            for para in cell.paragraphs:
                imgs = _find_images_in_element(para._element, rid_to_media)
                cell_images.extend(imgs)
            if cell_images:
                for img_name in cell_images:
                    img_id = _generate_image_id(doc_name, img_name)
                    image_registry[img_name] = img_id
                    cell_text += f" {{{{IMG:{img_id}}}}}"
            row_data.append(cell_text)
        rows.append(row_data)
    if not rows:
        return ""

    # Detect single-column tables containing code/diagrams
    if len(rows[0]) == 1:
        content = rows[0][0].replace('\\|', '|')
        # Check for Mermaid diagrams, SQL, code blocks
        code_indicators = [
            'sequenceDiagram', 'stateDiagram', 'erDiagram', 'flowchart',
            'classDiagram', 'gantt', 'pie', 'graph ',
            'SELECT ', 'INSERT ', 'CREATE TABLE', 'ALTER TABLE', 'DROP TABLE',
            'DELETE FROM', 'UPDATE ',
        ]
        lang_hints = {
            'sequenceDiagram': 'mermaid', 'stateDiagram': 'mermaid',
            'erDiagram': 'mermaid', 'flowchart': 'mermaid',
            'classDiagram': 'mermaid', 'gantt': 'mermaid', 'pie': 'mermaid',
            'graph ': 'mermaid',
            'SELECT ': 'sql', 'INSERT ': 'sql', 'CREATE TABLE': 'sql',
            'ALTER TABLE': 'sql', 'DROP TABLE': 'sql', 'DELETE FROM': 'sql',
            'UPDATE ': 'sql',
        }
        for indicator in code_indicators:
            if indicator in content:
                lang = lang_hints.get(indicator, '')
                # Restore newlines for readability
                raw_text = table.rows[0].cells[0].text.strip()
                return f"```{lang}\n{raw_text}\n```"

        # Check for Java/JSON/XML/Plaintext code patterns
        code_patterns = [
            ('Java ', 'java'), ('JSON ', 'json'), ('XML ', 'xml'),
            ('Plaintext ', 'text'), ('String ', 'java'),
            ('public ', 'java'), ('private ', 'java'),
            ('if (', 'java'), ('if(', 'java'),
        ]
        for pattern, lang in code_patterns:
            if content.startswith(pattern) or content.startswith(pattern.lower()):
                raw_text = table.rows[0].cells[0].text.strip()
                return f"```{lang}\n{raw_text}\n```"

    md_lines = []
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
    return "\n".join(md_lines)


def convert_docx_to_markdown(filepath):
    doc_name = os.path.splitext(os.path.basename(filepath))[0]
    rid_to_media, header_footer_media = _build_rid_to_media(filepath)
    ole_excel_map = _build_ole_excel_map(filepath)  # preview_image_name -> xlsx_zip_path
    ole_preview_names = set(ole_excel_map.keys())  # image names to skip (replaced by table)
    image_registry = {}
    image_data_map = {}
    md_lines = [f"# {doc_name}", ""]

    # Pre-open zip for xlsx parsing during body iteration
    _zip_handle = None
    try:
        _zip_handle = zipfile.ZipFile(filepath, 'r')
    except Exception:
        pass

    try:
        from docx import Document
        doc = Document(filepath)
    except ImportError:
        if _zip_handle:
            _zip_handle.close()
        return _convert_docx_raw(filepath)
    except Exception:
        if _zip_handle:
            _zip_handle.close()
        return _convert_docx_raw(filepath)

    # Build element-to-object lookup maps for O(1) access (avoids O(n²) inner loops)
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    # Iterate body elements in order to preserve table positions
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Find the corresponding paragraph object
            para = para_map.get(child)
            if para is None:
                continue
            text = para.text.strip()
            para_images = _find_images_in_element(para._element, rid_to_media)
            # Check if this paragraph is an OLE Excel embed with hint text
            _has_ole_embed = any(n in ole_preview_names for n in para_images) if para_images else False
            # Skip "点击图片可查看完整电子表格" hint text attached to OLE embeds
            if text and not (_has_ole_embed and '点击图片可查看完整电子表格' in text):
                level = _detect_heading_level(para)
                if level > 0:
                    md_lines.append("")
                    md_lines.append(f"{'#' * (level + 1)} {text}")
                    md_lines.append("")
                else:
                    md_lines.append(text)
                    md_lines.append("")
            if para_images:
                for img_name in para_images:
                    # Check if this image is a preview of an embedded Excel table
                    if img_name in ole_preview_names and _zip_handle:
                        xlsx_path = ole_excel_map[img_name]
                        table_md = _parse_embedded_xlsx(_zip_handle, xlsx_path)
                        if table_md:
                            md_lines.append("")
                            md_lines.append(f"<!-- 内嵌电子表格: {xlsx_path.split('/')[-1]} -->")
                            md_lines.append(table_md)
                            md_lines.append("<!-- /内嵌电子表格 -->")
                            md_lines.append("")
                            continue
                    img_id = _generate_image_id(doc_name, img_name)
                    image_registry[img_name] = img_id
                    md_lines.append(f"{{{{IMG:{img_id}}}}}")
                    md_lines.append("")
        elif tag == 'tbl':
            # Find the corresponding table object
            tbl = table_map.get(child)
            if tbl is not None:
                md_lines.append("")
                table_md = _table_to_markdown(tbl, rid_to_media, doc_name, image_registry)
                if table_md:
                    md_lines.append(table_md)
                    md_lines.append("")

    # Close the pre-opened zip handle (no longer needed after body iteration)
    if _zip_handle:
        _zip_handle.close()
        _zip_handle = None

    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            skipped_ids = []

            # Also detect orphan images (in zip but not referenced by any paragraph/table)
            referenced_names = set(image_registry.keys())
            for img_path in media:
                name = os.path.basename(img_path)
                # Skip OLE Excel preview images (already replaced with parsed table)
                if name in ole_preview_names:
                    continue
                # Skip watermark images (only referenced by headers/footers)
                if name in header_footer_media:
                    continue
                if name not in referenced_names:
                    # Orphan image: add to registry and markdown
                    ext = os.path.splitext(name)[1].lstrip('.')
                    if ext.lower() in ('emf', 'wmf'):
                        continue  # Skip vector-only orphans
                    img_data = z.read(img_path)
                    if len(img_data) < 500:
                        continue
                    try:
                        from PIL import Image as _Img
                        _tmp = _Img.open(BytesIO(img_data))
                        _w, _h = _tmp.size
                        if not _should_process_image(_w, _h):
                            continue
                    except Exception:
                        continue
                    img_id = _generate_image_id(doc_name, name)
                    image_registry[name] = img_id
                    image_data_map[img_id] = (img_data, ext)
                    md_lines.append(f"{{{{IMG:{img_id}}}}}")
                    md_lines.append("")

            for img_path in media:
                name = os.path.basename(img_path)
                # Skip OLE Excel preview images
                if name in ole_preview_names:
                    continue
                # Skip watermark images (only referenced by headers/footers)
                if name in header_footer_media:
                    continue
                if name in image_registry:
                    img_id = image_registry[name]
                    if img_id in image_data_map:
                        continue  # Already processed as orphan
                    ext = os.path.splitext(name)[1].lstrip('.')
                    img_data = z.read(img_path)

                    # Try to convert EMF/WMF to PNG via Pillow
                    if ext.lower() in ('emf', 'wmf'):
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            buf = BytesIO()
                            _tmp.save(buf, format='PNG')
                            img_data = buf.getvalue()
                            ext = 'png'
                        except Exception:
                            skipped_ids.append(img_id)
                            continue

                    if len(img_data) >= 500:
                        # Skip images that don't meet dimension threshold
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            _w, _h = _tmp.size
                            if not _should_process_image(_w, _h):
                                skipped_ids.append(img_id)
                                continue
                        except Exception:
                            pass
                        image_data_map[img_id] = (img_data, ext)
                    else:
                        skipped_ids.append(img_id)
            # Replace skipped image placeholders with annotation in markdown
            if skipped_ids:
                md_text = "\n".join(md_lines)
                for sid in skipped_ids:
                    placeholder = f"{{{{IMG:{sid}}}}}"
                    annotation = f"<!-- 已跳过(小图标/背景图): {sid} -->"
                    md_text = md_text.replace(placeholder, annotation)
                md_lines = md_text.split("\n")
    except Exception as e:
        sys.stderr.write(f"[MCP] Warning: error during image extraction: {e}\n")
        sys.stderr.flush()

    return "\n".join(md_lines), image_registry, image_data_map

def _convert_docx_raw(filepath):
    doc_name = os.path.splitext(os.path.basename(filepath))[0]
    rid_to_media, header_footer_media = _build_rid_to_media(filepath)
    ole_excel_map = _build_ole_excel_map(filepath)
    ole_preview_names = set(ole_excel_map.keys())
    image_registry = {}
    image_data_map = {}
    md_lines = [f"# {doc_name}", ""]

    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            if 'word/document.xml' in z.namelist():
                root = ET.fromstring(z.read('word/document.xml'))
                body = root.find(f'{W}body')
                if body is None:
                    body = root
                for child in body:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'p':
                        texts = [t.text for t in child.iter(f'{W}t') if t.text]
                        text = ''.join(texts).strip()
                        is_heading = False
                        para_images = _find_images_in_element(child, rid_to_media)
                        _has_ole_embed = any(n in ole_preview_names for n in para_images) if para_images else False
                        if text and not (_has_ole_embed and '点击图片可查看完整电子表格' in text):
                            pPr = child.find(f'{W}pPr')
                            if pPr is not None:
                                pStyle = pPr.find(f'{W}pStyle')
                                if pStyle is not None:
                                    sv = pStyle.get(f'{W}val', '')
                                    if 'Heading' in sv or 'heading' in sv:
                                        is_heading = True
                                        m = re.search(r'\d+', sv)
                                        level = int(m.group()) if m else 1
                                        md_lines.append("")
                                        md_lines.append(f"{'#' * (level + 1)} {text}")
                                        md_lines.append("")
                            if not is_heading:
                                md_lines.append(text)
                                md_lines.append("")
                        for img_name in para_images:
                            # Check if this is an OLE Excel preview image
                            if img_name in ole_preview_names:
                                xlsx_path = ole_excel_map[img_name]
                                table_md = _parse_embedded_xlsx(z, xlsx_path)
                                if table_md:
                                    md_lines.append("")
                                    md_lines.append(f"<!-- 内嵌电子表格: {xlsx_path.split('/')[-1]} -->")
                                    md_lines.append(table_md)
                                    md_lines.append("<!-- /内嵌电子表格 -->")
                                    md_lines.append("")
                                    continue
                            img_id = _generate_image_id(doc_name, img_name)
                            image_registry[img_name] = img_id
                            md_lines.append(f"{{{{IMG:{img_id}}}}}")
                            md_lines.append("")
                    elif tag == 'tbl':
                        rows = []
                        for tr in child.iter(f'{W}tr'):
                            row = []
                            for tc in tr.iter(f'{W}tc'):
                                cell_text = ''.join(t.text for t in tc.iter(f'{W}t') if t.text).strip()
                                cell_images = _find_images_in_element(tc, rid_to_media)
                                for img_name in cell_images:
                                    img_id = _generate_image_id(doc_name, img_name)
                                    image_registry[img_name] = img_id
                                    cell_text += f" {{{{IMG:{img_id}}}}}"
                                row.append(cell_text.replace('|', '\\|'))
                            if row:
                                rows.append(row)
                        if rows:
                            md_lines.append("")
                            md_lines.append("| " + " | ".join(rows[0]) + " |")
                            md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                            for row in rows[1:]:
                                while len(row) < len(rows[0]):
                                    row.append("")
                                md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
                            md_lines.append("")
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            skipped_ids = []

            # Detect orphan images (in zip but not referenced by any element)
            referenced_names = set(image_registry.keys())
            for img_path in media:
                name = os.path.basename(img_path)
                # Skip OLE Excel preview images
                if name in ole_preview_names:
                    continue
                # Skip watermark images (only referenced by headers/footers)
                if name in header_footer_media:
                    continue
                if name not in referenced_names:
                    ext = os.path.splitext(name)[1].lstrip('.')
                    if ext.lower() in ('emf', 'wmf'):
                        continue
                    img_data = z.read(img_path)
                    if len(img_data) < 500:
                        continue
                    try:
                        from PIL import Image as _Img
                        _tmp = _Img.open(BytesIO(img_data))
                        _w, _h = _tmp.size
                        if not _should_process_image(_w, _h):
                            continue
                    except Exception:
                        continue
                    img_id = _generate_image_id(doc_name, name)
                    image_registry[name] = img_id
                    image_data_map[img_id] = (img_data, ext)
                    md_lines.append(f"{{{{IMG:{img_id}}}}}")
                    md_lines.append("")

            for img_path in media:
                name = os.path.basename(img_path)
                # Skip OLE Excel preview images
                if name in ole_preview_names:
                    continue
                # Skip watermark images (only referenced by headers/footers)
                if name in header_footer_media:
                    continue
                if name in image_registry:
                    img_id = image_registry[name]
                    if img_id in image_data_map:
                        continue
                    ext = os.path.splitext(name)[1].lstrip('.')
                    img_data = z.read(img_path)

                    # Try to convert EMF/WMF to PNG
                    if ext.lower() in ('emf', 'wmf'):
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            buf = BytesIO()
                            _tmp.save(buf, format='PNG')
                            img_data = buf.getvalue()
                            ext = 'png'
                        except Exception:
                            skipped_ids.append(img_id)
                            continue

                    if len(img_data) >= 500:
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            _w, _h = _tmp.size
                            if not _should_process_image(_w, _h):
                                skipped_ids.append(img_id)
                                continue
                        except Exception:
                            pass
                        image_data_map[img_id] = (img_data, ext)
                    else:
                        skipped_ids.append(img_id)
            # Replace skipped image placeholders with annotation in markdown
            if skipped_ids:
                md_text = "\n".join(md_lines)
                for sid in skipped_ids:
                    placeholder = f"{{{{IMG:{sid}}}}}"
                    annotation = f"<!-- 已跳过(小图标/背景图): {sid} -->"
                    md_text = md_text.replace(placeholder, annotation)
                md_lines = md_text.split("\n")
    except Exception as e:
        md_lines.append(f"\n> 解析错误: {e}\n")

    return "\n".join(md_lines), image_registry, image_data_map

# ============================================================
# In-memory Store (backed by cache files)
# ============================================================

testcase_store = {
    "modules": [],
    "pending_images": [],
    "md_files": [],
}


def _sync_store_to_cache():
    """Persist critical store data to cache files."""
    _save_cache(CACHE_IMAGE_PROGRESS, {
        "pending_images": testcase_store["pending_images"],
        "md_files": testcase_store["md_files"],
    })
    # Always write testcases (even empty) to avoid stale cache
    _save_cache(CACHE_TESTCASES, {
        "modules": testcase_store["modules"],
    })


def _restore_store_from_cache():
    """Restore store from cache files (for session resume)."""
    img_data = _load_cache(CACHE_IMAGE_PROGRESS)
    if img_data:
        testcase_store["pending_images"] = img_data.get("pending_images", [])
        testcase_store["md_files"] = img_data.get("md_files", [])

    tc_data = _load_cache(CACHE_TESTCASES)
    if tc_data:
        testcase_store["modules"] = tc_data.get("modules", [])

# ============================================================
# Document Section Parser (for get_doc_summary / get_doc_section)
# ============================================================

def _parse_md_sections(md_content):
    """Parse markdown into sections by headings. Returns list of {heading, level, start, end, char_count}."""
    lines = md_content.split('\n')
    sections = []
    current = None
    in_code_block = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        # Track code blocks to avoid treating # comments as headings
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        if stripped.startswith('#') and not stripped.startswith('#!'):
            hashes = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            if not title:  # Skip lines that are just '#' with no title
                continue
            if current:
                current["end"] = i
                current["char_count"] = sum(len(lines[j]) for j in range(current["start"], i))
            current = {"heading": title, "level": hashes, "start": i, "end": len(lines), "char_count": 0}
            sections.append(current)

    if current:
        current["end"] = len(lines)
        current["char_count"] = sum(len(lines[j]) for j in range(current["start"], current["end"]))

    # If no headings found, treat entire content as one section
    if not sections:
        sections.append({
            "heading": "(全文)",
            "level": 1,
            "start": 0,
            "end": len(lines),
            "char_count": len(md_content)
        })

    return sections


def _build_doc_summary():
    """Build summary of all markdown docs: structure tree + stats."""
    md_files = testcase_store.get("md_files", [])
    summary = {"documents": [], "total_chars": 0, "total_sections": 0}

    for md_info in md_files:
        try:
            with open(md_info["path"], 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception:
            continue

        sections = _parse_md_sections(content)
        doc_entry = {
            "name": md_info["name"],
            "role": md_info.get("role", "primary"),
            "tag": md_info.get("tag"),
            "total_chars": len(content),
            "sections": []
        }
        for sec in sections:
            doc_entry["sections"].append({
                "heading": sec["heading"],
                "level": sec["level"],
                "char_count": sec["char_count"],
                "line_start": sec["start"],
                "line_end": sec["end"],
            })
        summary["documents"].append(doc_entry)
        summary["total_chars"] += len(content)
        summary["total_sections"] += len(sections)

    # Also save to cache
    _save_cache(CACHE_DOC_SUMMARY, summary)
    return summary

# ============================================================
# XMind Export
# ============================================================

def _esc(text):
    if not text:
        return ""
    return (text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            .replace('"', '&quot;').replace("'", '&apos;'))


def create_xmind_file(modules, output_path):
    counter = [0]
    def nid():
        counter[0] += 1
        return f"t{counter[0]}"
    def topic(title, children=None):
        x = f'<topic id="{nid()}"><title>{_esc(title)}</title>'
        if children:
            x += '<children><topics type="attached">' + ''.join(children) + '</topics></children>'
        return x + '</topic>'

    mod_topics = []
    for m in modules:
        sub_topics = []
        for s in m.get("sub_modules", []):
            case_topics = []
            for c in s.get("test_cases", []):
                # Build chain: 用例标题 → 前置条件 → 执行步骤 → 预期结果
                inner = None
                if c.get("expected_result"):
                    inner = topic(f"预期结果: {c['expected_result']}")
                steps = c.get("steps", [])
                if steps:
                    if isinstance(steps, str):
                        steps_text = steps
                    else:
                        steps_text = "\n".join(f"{i}. {step}" for i, step in enumerate(steps, 1))
                    inner = topic(f"执行步骤:\n{steps_text}", [inner] if inner else None)
                if c.get("preconditions"):
                    inner = topic(f"前置条件: {c['preconditions']}", [inner] if inner else None)
                case_topics.append(topic(c.get("title", "未命名用例"), [inner] if inner else None))
            sub_topics.append(topic(s.get("name", ""), case_topics))
        mod_topics.append(topic(m.get("name", ""), sub_topics))

    content_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<xmap-content xmlns="urn:xmind:xmap:xmlns:content:2.0"
  xmlns:fo="http://www.w3.org/1999/XSL/Format"
  xmlns:svg="http://www.w3.org/2000/svg"
  xmlns:xhtml="http://www.w3.org/1999/xhtml"
  xmlns:xlink="http://www.w3.org/1999/xlink" version="2.0">
  <sheet id="sheet_1"><title>测试用例</title>{topic("测试用例", mod_topics)}</sheet>
</xmap-content>'''

    manifest_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<manifest xmlns="urn:xmind:xmap:xmlns:manifest:1.0">
  <file-entry full-path="content.xml" media-type="text/xml"/>
  <file-entry full-path="META-INF/" media-type=""/>
  <file-entry full-path="META-INF/manifest.xml" media-type="text/xml"/>
</manifest>'''

    meta_xml = ('<?xml version="1.0" encoding="UTF-8"?>'
                '<meta xmlns="urn:xmind:xmap:xmlns:meta:2.0" version="2.0">'
                '<Author><Name>TestCase Generator</Name></Author></meta>')

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('content.xml', content_xml)
        zf.writestr('META-INF/manifest.xml', manifest_xml)
        zf.writestr('meta.xml', meta_xml)
    return output_path

# ============================================================
# MCP Tool Definitions
# ============================================================

TOOLS = [
    {
        "name": "setup_environment",
        "description": "启动检查: 1) 检查并安装Python依赖 2) 检查并创建工作目录(doc/, .tmp/doc_mk/, .tmp/picture/, .tmp/cache/) 3) 检测缓存任务。如检测到缓存任务，返回 has_cache=true 和缓存详情，agent 需询问用户是继续上次任务还是开始新任务。",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "clear_cache",
        "description": "清除所有缓存任务数据(.tmp/cache/下的所有状态文件、.tmp/doc_mk/下的markdown文件、.tmp/picture/下的图片)，用于开始全新的用例生成任务。调用前应先确认用户意图。",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "parse_documents",
        "description": "Parse .docx files from doc/ directory: convert to markdown (.tmp/doc_mk/), extract images to .tmp/picture/. Returns file list and pending image count. State is persisted to .tmp/cache/ for cross-session resume. Will block if there's an in-progress workflow (pass force=true to override). 支持文档打标分类：文件名含【主prd】【主概设】【主后端概设】【主前端概设】的为主文档（提取图片+生成用例），其余为辅助资料（仅解析文字，按需查阅）。如果所有文档都没有标签，则全部视为主文档。",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Directory containing .docx files (default: cwd)"},
                "file_patterns": {"type": "string", "description": "Glob pattern (default: *.docx)"},
                "force": {"type": "boolean", "description": "Force re-parse even if there's in-progress work (default: false)"}
            },
            "required": []
        }
    },

    {
        "name": "get_workflow_state",
        "description": "Get current workflow state for session resume. Returns phase progress, pending work, and resume instructions. Call this at the start of a new session to continue previous work.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Workspace directory (default: cwd). Used to locate .tmp/cache/."}
            },
            "required": []
        }
    },
    {
        "name": "get_doc_summary",
        "description": "Get document structure summary (heading tree + char counts per section) without loading full content. Use this to plan which sections to read with get_doc_section. 文档按角色分类显示：📌主文档（用例生成目标）和📎辅助资料（按需查阅补充用例设计）。",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "get_doc_section",
        "description": "Read a specific section of a markdown document by heading name. Returns only that section's content, reducing context usage vs loading the full document.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "doc_name": {"type": "string", "description": "Markdown filename (e.g. 'xxx.md')"},
                "section_heading": {"type": "string", "description": "Heading text to match (fuzzy match supported)"},
                "include_subsections": {"type": "boolean", "description": "Include child sections (default: true)"}
            },
            "required": ["doc_name"]
        }
    },
    {
        "name": "get_parsed_markdown",
        "description": "Read all processed markdown files. WARNING: may be large. Prefer get_doc_summary + get_doc_section for large documents to avoid context overflow.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "save_testcases",
        "description": "Save test cases. Also persists to .tmp/cache/testcases.json for cross-session access. Supports incremental save via append_module parameter.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "modules": {"type": "array", "description": "Test case module list", "items": {"type": "object"}},
                "append_module": {"type": "object", "description": "Single module to append to existing cases (for incremental generation)"}
            },
            "required": []
        }
    },
    {
        "name": "get_testcases",
        "description": "Get all current test cases. Loads from cache if memory is empty.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "export_xmind",
        "description": "Export test cases to XMind format. File named as 需求名_testCase.xmind by default.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "output_path": {"type": "string", "description": "Output file path (default: 需求名_testCase.xmind)"},
                "requirement_name": {"type": "string", "description": "Requirement name for file naming (auto-detected from docs if not provided)"}
            },
            "required": []
        }
    },
    {
        "name": "review_module_structure",
        "description": "Review test case module structure for balance, duplicates, empty modules, and quality issues. Call this after initial generation and before final review to optimize module organization.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "export_report",
        "description": "Generate test case report as markdown file (需求名_testCaseReport.md). Includes module overview, coverage dimensions, and requirement questions.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "requirement_name": {"type": "string", "description": "Requirement name for file naming (auto-detected if not provided)"},
                "output_dir": {"type": "string", "description": "Output directory (default: workspace root)"},
                "questions": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of requirement questions/confirmation points discovered during analysis"
                }
            },
            "required": []
        }
    },
    {
        "name": "configure_llm_api",
        "description": "打开GUI窗口配置外部多模态LLM API，用于图片解析。支持配置API地址、API Key、测试连接、获取模型列表、选择模型、多线程设置。配置会自动保存，下次打开时恢复上次输入。用户选择'图片+文本解析'模式时调用此工具。",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "process_images_with_llm",
        "description": "使用已配置的外部多模态LLM API批量处理所有待处理图片。支持多线程并发。需先调用configure_llm_api配置API。处理完成后自动将分析结果回填到Markdown文档对应位置。",
        "inputSchema": {
            "type": "object",
            "properties": {
                "force_reprocess": {"type": "boolean", "description": "强制重新处理已处理的图片 (default: false)"}
            },
            "required": []
        }
    }
]

# ============================================================
# Tool Handlers
# ============================================================


def handle_setup_environment(args):
    results = []
    all_ok = True
    results.append(f"Python {sys.version.split()[0]}")

    # 1. Check Python dependencies
    deps = {"docx": "python-docx", "PIL": "Pillow", "openpyxl": "openpyxl"}
    for imp, pip_name in deps.items():
        try:
            __import__(imp)
            results.append(f"  [ok] {pip_name}")
        except ImportError:
            results.append(f"  [installing] {pip_name}...")
            if _ensure_pkg(imp, pip_name):
                results.append(f"  [ok] {pip_name} installed")
            else:
                results.append(f"  [FAIL] {pip_name}")
                all_ok = False

    # 2. Ensure working directories exist
    workspace = _workspace()
    dirs_to_check = {
        "doc": os.path.join(workspace, "doc"),
        ".tmp/doc_mk": TMP_DOC_DIR,
        ".tmp/picture": TMP_PIC_DIR,
        ".tmp/cache": TMP_CACHE_DIR,
    }
    results.append("")
    results.append("工作目录检查:")
    for label, dir_path in dirs_to_check.items():
        if os.path.isdir(dir_path):
            results.append(f"  [ok] {label}/")
        else:
            try:
                os.makedirs(dir_path, exist_ok=True)
                results.append(f"  [created] {label}/")
            except Exception as e:
                results.append(f"  [FAIL] {label}/ — {e}")
                all_ok = False

    # 3. Check for cached tasks
    has_cache = False
    cache_info = {}
    existing_state = _load_cache(CACHE_PHASE_STATE)
    if existing_state and existing_state.get("phases"):
        has_cache = True
        # Gather cache details
        _restore_store_from_cache()
        pending = testcase_store.get("pending_images", [])
        total_imgs = len(pending)
        processed_imgs = sum(1 for p in pending if p["processed"])
        skipped_imgs = sum(1 for p in pending if p.get("skipped"))
        modules = testcase_store.get("modules", [])
        total_cases = sum(len(s.get("test_cases", [])) for m in modules for s in m.get("sub_modules", []))
        current_phase = existing_state.get("current_phase", "unknown")

        cache_info = {
            "current_phase": current_phase,
            "total_images": total_imgs,
            "processed_images": processed_imgs,
            "skipped_images": skipped_imgs,
            "unprocessed_images": total_imgs - processed_imgs,
            "module_count": len(modules),
            "total_cases": total_cases,
        }
        results.append("")
        results.append("⚠️ 检测到缓存任务:")
        results.append(f"  当前阶段: {current_phase}")
        if total_imgs > 0:
            skip_note = f" (其中 {skipped_imgs} 张因清晰度跳过)" if skipped_imgs > 0 else ""
            results.append(f"  图片处理: {processed_imgs}/{total_imgs}{skip_note}")
        if modules:
            results.append(f"  已生成用例: {len(modules)} 模块, {total_cases} 用例")
        results.append("")
        results.append("请询问用户:")
        results.append("  1. 继续上次任务 — 调用 get_workflow_state 恢复进度")
        results.append("  2. 开始新任务 — 调用 clear_cache 清除缓存后开始新的用例生成")

    results.append("")
    results.append("OK - environment ready" if all_ok else "WARN - some deps failed")
    return {
        "content": [{"type": "text", "text": "\n".join(results)}],
        "all_ok": all_ok,
        "has_cache": has_cache,
        "cache_info": cache_info,
    }




def handle_clear_cache(args):
    """Clear all cached task data for a fresh start."""
    cleared = []

    # Clear cache files
    for cache_file in (CACHE_PHASE_STATE, CACHE_IMAGE_PROGRESS, CACHE_TESTCASES, CACHE_DOC_SUMMARY):
        cache_fp = _cache_path(cache_file)
        if os.path.exists(cache_fp):
            os.remove(cache_fp)
            cleared.append(cache_file)

    # Clear generated markdown files
    if os.path.isdir(TMP_DOC_DIR):
        shutil.rmtree(TMP_DOC_DIR)
        os.makedirs(TMP_DOC_DIR, exist_ok=True)
        cleared.append(".tmp/doc_mk/*")

    # Clear extracted images
    if os.path.isdir(TMP_PIC_DIR):
        shutil.rmtree(TMP_PIC_DIR)
        os.makedirs(TMP_PIC_DIR, exist_ok=True)
        cleared.append(".tmp/picture/*")

    # Reset in-memory store
    testcase_store["modules"] = []
    testcase_store["pending_images"] = []
    testcase_store["md_files"] = []

    msg = "✓ 缓存已清除:\n  " + "\n  ".join(cleared) if cleared else "没有需要清除的缓存。"
    msg += "\n\n可以开始新的用例生成任务了。请确认文档已放入 doc/ 目录，然后调用 parse_documents 开始。"
    return {"content": [{"type": "text", "text": msg}]}


def handle_parse_documents(args):
    directory = args.get("directory", _workspace())
    pattern = args.get("file_patterns", "*.docx")
    force = args.get("force", False)
    _update_workspace(directory)

    # Default: look in doc/ subdirectory
    doc_dir = os.path.join(directory, "doc")
    search_dir = doc_dir if os.path.isdir(doc_dir) else directory

    # Protection: check if there's an in-progress workflow
    if not force:
        existing_state = _load_cache(CACHE_PHASE_STATE)
        if existing_state:
            phases = existing_state.get("phases", {})
            # Check for in-progress image analysis
            img_phase = phases.get("image_analysis", {})
            if img_phase.get("status") == "in_progress":
                processed = img_phase.get("processed", 0)
                total = img_phase.get("total", 0)
                return {"content": [{"type": "text", "text": (
                    f"⚠️ 检测到未完成的图片处理进度 ({processed}/{total})。\n"
                    f"重新解析会丢失已有进度。如需继续处理，请调用 get_workflow_state 恢复。\n"
                    f"如确认要重新开始，请传入 force=true 参数。"
                )}], "blocked": True}
            # Check for existing test cases that would be lost
            gen_phase = phases.get("generation", {})
            if gen_phase.get("status") in ("in_progress", "completed"):
                module_count = gen_phase.get("module_count", 0)
                if module_count > 0:
                    return {"content": [{"type": "text", "text": (
                        f"⚠️ 检测到已生成的测试用例 ({module_count} 个模块)。\n"
                        f"重新解析会丢失所有已生成的用例。如需继续当前任务，请调用 get_workflow_state 恢复。\n"
                        f"如确认要重新开始，请传入 force=true 参数。"
                    )}], "blocked": True}

    files = glob.glob(os.path.join(search_dir, "**", pattern), recursive=True)
    if not files:
        files = glob.glob(os.path.join(search_dir, pattern))
    if not files:
        return {"content": [{"type": "text", "text": f"No .docx files found in {search_dir}"}]}

    # Clean doc/picture dirs and reset all cache state for fresh start
    for d in (TMP_DOC_DIR, TMP_PIC_DIR):
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d, exist_ok=True)
    os.makedirs(TMP_CACHE_DIR, exist_ok=True)

    # Reset all cache state files when re-parsing
    for cache_file in (CACHE_PHASE_STATE, CACHE_IMAGE_PROGRESS, CACHE_TESTCASES, CACHE_DOC_SUMMARY):
        cache_fp = _cache_path(cache_file)
        if os.path.exists(cache_fp):
            os.remove(cache_fp)

    # Reset in-memory store
    testcase_store["modules"] = []
    testcase_store["pending_images"] = []
    testcase_store["md_files"] = []

    # Initialize clean phase state
    _reset_phase_state()

    all_md_files = []
    all_pending_images = []
    content_parts = []

    # Phase 1: Classify all documents
    doc_classifications = []
    for fpath in files:
        basename = os.path.basename(fpath)
        role, tag = _classify_document(basename)
        doc_classifications.append((fpath, role, tag))

    # If no documents have primary tags, treat ALL as primary (backward compat)
    has_any_primary = any(role == 'primary' for _, role, _ in doc_classifications)
    if not has_any_primary:
        doc_classifications = [(fp, 'primary', None) for fp, _, _ in doc_classifications]

    primary_count = sum(1 for _, r, _ in doc_classifications if r == 'primary')
    ref_count = sum(1 for _, r, _ in doc_classifications if r == 'reference')

    for fpath, doc_role, doc_tag in doc_classifications:
        try:
            doc_name = os.path.splitext(os.path.basename(fpath))[0]
            md_text, image_registry, image_data_map = convert_docx_to_markdown(fpath)

            md_filename = f"{doc_name}.md"
            md_path = os.path.join(TMP_DOC_DIR, md_filename)
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            all_md_files.append({
                "name": md_filename, "path": md_path,
                "role": doc_role, "tag": doc_tag
            })

            # Only extract images for primary documents
            if doc_role == 'primary':
                skipped = 0
                for img_id, (img_data, ext) in image_data_map.items():
                    resized_data, mime = _resize_image(img_data, ext)
                    out_ext = ".png" if "png" in mime else ".jpg"
                    img_filename = os.path.splitext(img_id)[0] + out_ext
                    img_path = os.path.join(TMP_PIC_DIR, img_filename)
                    # Store relative path for cross-workspace portability
                    img_rel_path = os.path.join(".tmp", "picture", img_filename)
                    with open(img_path, 'wb') as f:
                        f.write(resized_data)
                    all_pending_images.append({
                        "id": img_id, "filename": img_filename,
                        "path": img_path, "rel_path": img_rel_path,
                        "mime": mime, "size": len(img_data), "source_doc": doc_name,
                        "processed": False
                    })

                for orig_name, uid in image_registry.items():
                    if uid not in image_data_map:
                        skipped += 1

                role_label = f"[主文档{doc_tag or ''}]" if doc_tag else "[主文档]"
                content_parts.append({
                    "type": "text",
                    "text": (f"✓ {role_label} {os.path.basename(fpath)} → {md_filename}\n"
                             f"  图片: {len(image_data_map)} 张提取, {skipped} 张跳过(EMF/WMF/过小)")
                })
            else:
                # Reference document: text only, no image processing
                content_parts.append({
                    "type": "text",
                    "text": (f"✓ [辅助资料] {os.path.basename(fpath)} → {md_filename}\n"
                             f"  仅解析文字内容，图片不处理（{len(image_data_map)} 张跳过）")
                })
        except Exception as e:
            content_parts.append({"type": "text", "text": f"✗ Error parsing {fpath}: {e}"})

    testcase_store["pending_images"] = all_pending_images
    testcase_store["md_files"] = all_md_files

    # Persist to cache
    _save_phase_state("parse", "completed", {
        "file_count": len(all_md_files),
        "total_images": len(all_pending_images)
    })
    _sync_store_to_cache()

    total_imgs = len(all_pending_images)
    summary = (f"\n转换完成:\n"
               f"  Markdown 文件: {len(all_md_files)} 个 → {TMP_DOC_DIR}\n")
    if ref_count > 0:
        summary += f"    主文档: {primary_count} 个（生成用例+处理图片）\n"
        summary += f"    辅助资料: {ref_count} 个（仅文字，按需查阅）\n"
    summary += (f"  图片文件: {total_imgs} 张 → {TMP_PIC_DIR}\n"
                f"  缓存目录: {TMP_CACHE_DIR}\n")
    if total_imgs > 0:
        summary += ("\n图片处理：\n"
                    "  调用 configure_llm_api 配置外部多模态LLM API，然后调用 process_images_with_llm 批量处理\n"
                    "全部处理完成后调用 get_doc_summary 获取文档结构概览。")
    else:
        summary += "\n无需处理图片，可直接调用 get_doc_summary 获取文档结构概览。"

    content_parts.append({"type": "text", "text": summary})
    return {"content": content_parts, "file_count": len(all_md_files), "total_images": total_imgs}

IMAGE_ANALYSIS_PROMPT = (
    "你是一位资深测试开发专家，正在从需求文档中提取测试用例设计所需的信息。\n"
    "请先判断这张图片属于以下哪种类型，然后按对应规则提取具体内容：\n\n"
    "1. 数据表/字段定义 → 用 markdown 表格逐行提取每个字段的：字段名、数据类型、长度、是否必填、默认值、描述。不要遗漏任何一行。\n"
    "2. 流程图/状态机 → 列出所有节点和转换条件，用 A --[条件]--> B 格式描述每条路径（含异常分支），确保无遗漏。\n"
    "3. ER图/架构图 → 列出所有实体及属性，标注实体间关系（一对多等）和外键。\n"
    "4. UI界面/原型图 → 列出所有表单字段（标签、输入类型、可选值）、按钮、表格列头及示例数据。\n"
    "5. 接口/参数定义 → 用 markdown 表格逐个提取参数名、类型、是否必填、取值范围、描述。\n"
    "6. 其他 → 提取所有可见文字和关键信息。\n\n"
    "【输出格式】先用一行标注图片类型，然后输出提取的具体内容。\n"
    "【核心原则】只提取具体数据，禁止笼统概括。看到表格就逐行抄录，看到流程图就逐条列出路径。\n"
    "【测试视角】重点关注：字段约束（长度、格式、必填）、状态转换条件、边界值、业务规则，这些是设计测试用例的关键依据。\n"
    "【无法识别】如果图片模糊、分辨率过低、内容不清晰导致无法准确提取具体信息，"
    "请直接输出一行：`[UNREADABLE]` 加上简短原因说明（如 `[UNREADABLE] 图片模糊，文字无法辨认`）。"
    "不要猜测或编造内容。"
)

# Marker returned by LLM when image is unreadable
_UNREADABLE_MARKER = "[UNREADABLE]"




def handle_get_workflow_state(args):
    """Return current workflow state for session resume."""
    directory = args.get("directory", _workspace())
    _update_workspace(directory)

    # Try to restore from cache
    _restore_store_from_cache()

    state = _load_cache(CACHE_PHASE_STATE)
    if not state:
        return {"content": [{"type": "text", "text": "没有找到已保存的工作流状态。请从 parse_documents 开始新的工作流。"}],
                "has_state": False}

    pending = testcase_store.get("pending_images", [])
    total_imgs = len(pending)
    processed_imgs = sum(1 for p in pending if p["processed"])
    skipped_imgs = sum(1 for p in pending if p.get("skipped"))
    unprocessed_imgs = total_imgs - processed_imgs

    md_files = testcase_store.get("md_files", [])
    modules = testcase_store.get("modules", [])
    total_cases = sum(len(s.get("test_cases", [])) for m in modules for s in m.get("sub_modules", []))

    # Determine current phase and next action
    phases = state.get("phases", {})
    lines = ["📋 工作流状态恢复:", ""]

    # Parse phase
    parse_status = phases.get("parse", {}).get("status", "pending")
    lines.append(f"  阶段1 文档解析: {parse_status}")
    if md_files:
        lines.append(f"    - {len(md_files)} 个 Markdown 文件")

    # Image analysis phase — auto-fix stale status
    img_status = phases.get("image_analysis", {}).get("status", "pending")
    if img_status == "in_progress" and total_imgs > 0 and unprocessed_imgs == 0:
        # All images processed but status not updated — fix it
        img_status = "completed"
        _save_phase_state("image_analysis", "completed")
    lines.append(f"  阶段2 图片分析: {img_status}")
    if total_imgs > 0:
        skip_note = f", {skipped_imgs} 张跳过(不清晰)" if skipped_imgs > 0 else ""
        lines.append(f"    - {processed_imgs}/{total_imgs} 张已处理{skip_note}, {unprocessed_imgs} 张待处理")

    # Generation phase
    gen_status = phases.get("generation", {}).get("status", "pending")
    lines.append(f"  阶段3 用例生成: {gen_status}")
    if modules:
        lines.append(f"    - {len(modules)} 个模块, {total_cases} 个用例")

    # Export phase
    export_status = phases.get("export", {}).get("status", "pending")
    lines.append(f"  阶段4 导出: {export_status}")

    # Determine resume instruction
    lines.append("")
    img_completed = img_status == "completed"
    has_unprocessed_images = unprocessed_imgs > 0
    has_testcases = total_cases > 0

    if has_unprocessed_images:
        lines.append(f"▶ 继续操作: 有 {unprocessed_imgs} 张图片待处理")
        lines.append("  - 调用 configure_llm_api + process_images_with_llm 使用外部LLM批量处理")
    elif not img_completed and total_imgs > 0:
        lines.append("▶ 继续操作: 调用 process_images_with_llm 检查图片处理状态")
    elif (img_completed or total_imgs == 0) and not has_testcases:
        lines.append("▶ 继续操作: 调用 get_doc_summary 获取文档结构，然后按模块生成测试用例")
    elif has_testcases and export_status != "completed":
        lines.append("▶ 继续操作: 调用 get_testcases 查看已有用例，可继续生成或调用 review_module_structure 审查模块结构，最后调用 export_xmind 和 export_report 导出")
    elif export_status == "completed":
        lines.append("▶ 已导出 XMind 和测试报告。可以继续修改用例并重新导出，或确认用例完善后结束流程。")
        lines.append("  - 如需修改用例: 调用 get_testcases 查看当前用例，修改后调用 save_testcases(append_module=...) 保存，再调用 export_xmind + export_report 重新导出")
        lines.append("  - 如需重新生成: 调用 clear_cache 清除缓存后重新开始")
    elif parse_status == "completed":
        lines.append("▶ 继续操作: 调用 configure_llm_api + process_images_with_llm 开始处理图片")
    else:
        lines.append("▶ 继续操作: 调用 parse_documents 开始新的工作流")

    return {
        "content": [{"type": "text", "text": "\n".join(lines)}],
        "has_state": True,
        "current_phase": state.get("current_phase"),
        "unprocessed_images": unprocessed_imgs,
        "total_cases": total_cases,
        "module_count": len(modules),
    }

def handle_get_doc_summary(args):
    """Return document structure summary without full content."""
    if not testcase_store.get("md_files"):
        _restore_store_from_cache()
    if not testcase_store.get("md_files"):
        return {"content": [{"type": "text", "text": "No markdown files found. Call parse_documents first."}]}

    summary = _build_doc_summary()
    lines = ["📄 文档结构概览:", ""]

    primary_docs = []
    ref_docs = []
    for doc in summary["documents"]:
        role = doc.get("role", "primary")
        if role == "primary":
            primary_docs.append(doc)
        else:
            ref_docs.append(doc)

    if primary_docs:
        lines.append("📌 主文档（用例生成目标）:")
        for doc in primary_docs:
            tag = doc.get("tag") or ""
            lines.append(f"  📁 {tag} {doc['name']} ({doc['total_chars']} 字符)")
            for sec in doc["sections"]:
                indent = "    " + "  " * sec["level"]
                lines.append(f"{indent}{'#' * sec['level']} {sec['heading']} ({sec['char_count']} 字符)")
            lines.append("")

    if ref_docs:
        lines.append("📎 辅助资料（按需查阅）:")
        for doc in ref_docs:
            lines.append(f"  📁 {doc['name']} ({doc['total_chars']} 字符)")
            for sec in doc["sections"]:
                indent = "    " + "  " * sec["level"]
                lines.append(f"{indent}{'#' * sec['level']} {sec['heading']} ({sec['char_count']} 字符)")
            lines.append("")

    lines.append(f"总计: {len(summary['documents'])} 个文档 ({len(primary_docs)} 主文档 + {len(ref_docs)} 辅助资料), {summary['total_sections']} 个章节, {summary['total_chars']} 字符")
    lines.append("")
    if primary_docs:
        lines.append("请优先按主文档的模块调用 get_doc_section(doc_name, section_heading) 分段读取内容，")
        lines.append("每读取一个模块就生成该模块的测试用例。辅助资料在需要补充信息时按需查阅。")
    else:
        lines.append("请按模块调用 get_doc_section(doc_name, section_heading) 分段读取内容，")
        lines.append("每读取一个模块就生成该模块的测试用例，避免一次性加载全部文档。")

    return {
        "content": [{"type": "text", "text": "\n".join(lines)}],
        "summary": summary,
    }


def handle_get_doc_section(args):
    """Read a specific section from a markdown document."""
    doc_name = args.get("doc_name", "")
    section_heading = args.get("section_heading", "")
    include_sub = args.get("include_subsections", True)

    if not doc_name:
        return {"content": [{"type": "text", "text": "Missing required parameter: doc_name"}]}

    if not testcase_store.get("md_files"):
        _restore_store_from_cache()

    # Find the markdown file
    md_path = None
    for md_info in testcase_store.get("md_files", []):
        if md_info["name"] == doc_name:
            md_path = md_info["path"]
            break

    if not md_path:
        # Try fuzzy match
        for md_info in testcase_store.get("md_files", []):
            if doc_name in md_info["name"] or md_info["name"] in doc_name:
                md_path = md_info["path"]
                doc_name = md_info["name"]
                break

    if not md_path or not os.path.exists(md_path):
        available = [m["name"] for m in testcase_store.get("md_files", [])]
        return {"content": [{"type": "text", "text": f"Document not found: {doc_name}\nAvailable: {', '.join(available)}"}]}

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Error reading {md_path}: {e}"}]}

    # If no section specified, return full content
    if not section_heading:
        return {"content": [{"type": "text", "text": f"FILE: {doc_name}\n\n{content}"}]}

    lines = content.split('\n')
    sections = _parse_md_sections(content)

    # Find matching section (fuzzy)
    target_sec = None
    for sec in sections:
        if section_heading in sec["heading"] or sec["heading"] in section_heading:
            target_sec = sec
            break
    # Looser match
    if not target_sec:
        section_lower = section_heading.lower()
        for sec in sections:
            if section_lower in sec["heading"].lower() or sec["heading"].lower() in section_lower:
                target_sec = sec
                break

    if not target_sec:
        available_headings = [s["heading"] for s in sections]
        return {"content": [{"type": "text", "text": (
            f"Section not found: '{section_heading}' in {doc_name}\n"
            f"Available sections:\n" + "\n".join(f"  - {h}" for h in available_headings)
        )}]}

    # Determine end line
    start = target_sec["start"]
    if include_sub:
        # Include all child sections (find next section at same or higher level)
        end = len(lines)
        for sec in sections:
            if sec["start"] > start and sec["level"] <= target_sec["level"]:
                end = sec["start"]
                break
    else:
        # Only this section's own content (up to next heading of any level)
        end = target_sec["end"]

    section_content = '\n'.join(lines[start:end])

    return {
        "content": [{"type": "text", "text": f"SECTION: {target_sec['heading']} (from {doc_name})\n\n{section_content}"}],
        "section_heading": target_sec["heading"],
        "char_count": len(section_content),
    }

def handle_get_parsed_markdown(args):
    """Read all markdown files - full content. Prefer get_doc_summary + get_doc_section for large docs."""
    if not testcase_store.get("md_files"):
        _restore_store_from_cache()

    md_files = testcase_store.get("md_files", [])
    if not md_files:
        return {"content": [{"type": "text", "text": "No markdown files found. Call parse_documents first."}]}

    pending = testcase_store.get("pending_images", [])
    unprocessed = sum(1 for p in pending if not p["processed"])

    content_parts = []
    total_chars = 0

    for md_info in md_files:
        try:
            with open(md_info["path"], 'r', encoding='utf-8') as f:
                md_content = f.read()
            total_chars += len(md_content)
            content_parts.append({
                "type": "text",
                "text": f"\n{'='*60}\nFILE: {md_info['name']}\n{'='*60}\n\n{md_content}"
            })
        except Exception as e:
            content_parts.append({"type": "text", "text": f"Error reading {md_info['path']}: {e}"})

    summary = f"\n共 {len(md_files)} 个文档已加载 ({total_chars} 字符)。"
    if unprocessed > 0:
        summary += f"\n注意: 还有 {unprocessed} 张图片未处理。"
    if total_chars > 30000:
        summary += "\n⚠️ 文档内容较大，建议使用 get_doc_summary + get_doc_section 分段读取以减少上下文占用。"

    content_parts.append({"type": "text", "text": summary})
    return {"content": content_parts, "file_count": len(md_files), "total_chars": total_chars}


def handle_save_testcases(args):
    modules = args.get("modules", None)
    append_module = args.get("append_module", None)

    if append_module:
        # Validate structure
        if not isinstance(append_module, dict):
            return {"content": [{"type": "text", "text": "Error: append_module must be a JSON object, not an array or primitive."}]}
        if "name" not in append_module:
            return {"content": [{"type": "text", "text": "Error: append_module must have a 'name' field."}]}
        if "sub_modules" not in append_module:
            append_module["sub_modules"] = []

        # Incremental: append one module (replace if same name exists)
        if not testcase_store["modules"]:
            _restore_store_from_cache()
        # Replace existing module with same name, or append new
        mod_name = append_module.get("name", "")
        replaced = False
        for i, existing in enumerate(testcase_store["modules"]):
            if existing.get("name") == mod_name:
                testcase_store["modules"][i] = append_module
                replaced = True
                break
        if not replaced:
            testcase_store["modules"].append(append_module)
        _sync_store_to_cache()
        _save_phase_state("generation", "in_progress", {
            "module_count": len(testcase_store["modules"])
        })
        total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"] for s in m.get("sub_modules", []))
        action = "替换" if replaced else "追加"
        return {"content": [{"type": "text", "text": (
            f"✓ {action}模块 '{mod_name}', "
            f"当前共 {len(testcase_store['modules'])} 个模块, {total} 个用例"
        )}]}

    if modules is None:
        return {"content": [{"type": "text", "text": (
            "Missing parameter: modules or append_module.\n"
            "必须提供 modules（全量数组）或 append_module（单个模块对象）之一。\n\n"
            "⚠️ 如果你正在尝试全量替换但数据量太大导致参数丢失，请改用 append_module 逐个模块更新：\n"
            "  1. 调用 get_testcases 获取当前用例\n"
            "  2. 对需要修改的模块，逐个调用 save_testcases(append_module={修改后的单个模块对象})\n"
            "  3. append_module 会自动按模块名替换已有模块\n"
            "  4. 不需要修改的模块无需重新提交"
        )}]}

    if not isinstance(modules, list):
        return {"content": [{"type": "text", "text": "Error: modules must be a JSON array."}]}

    # Validate each module has required fields
    for i, m in enumerate(modules):
        if not isinstance(m, dict):
            return {"content": [{"type": "text", "text": f"Error: modules[{i}] must be a JSON object."}]}
        if "name" not in m:
            return {"content": [{"type": "text", "text": f"Error: modules[{i}] must have a 'name' field."}]}
        if "sub_modules" not in m:
            m["sub_modules"] = []

    testcase_store["modules"] = modules
    _sync_store_to_cache()
    _save_phase_state("generation", "completed", {"module_count": len(modules)})

    total = sum(len(s.get("test_cases", [])) for m in modules for s in m.get("sub_modules", []))
    return {"content": [{"type": "text", "text": f"Saved {len(modules)} modules, {total} test cases. 已持久化到 .tmp/cache/testcases.json"}]}


def handle_get_testcases(args):
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"]
                for s in m.get("sub_modules", []))
    return {
        "content": [{"type": "text", "text": json.dumps(testcase_store["modules"], ensure_ascii=False, indent=2)}],
        "module_count": len(testcase_store["modules"]),
        "total_cases": total,
    }


def _get_requirement_name():
    """Extract requirement name from parsed documents for file naming."""
    md_files = testcase_store.get("md_files", [])
    if not md_files:
        _restore_store_from_cache()
        md_files = testcase_store.get("md_files", [])
    # Use the first requirement doc name (strip common prefixes)
    for md_info in md_files:
        name = md_info.get("name", "")
        name = os.path.splitext(name)[0]
        # Prefer requirement docs over design docs
        if "需求" in name or "requirement" in name.lower():
            # Clean up common prefixes like [需求] or 【主prd】
            name = re.sub(r'^[\[【][^\]】]*[\]】]', '', name).strip()
            if name:
                return name
    # Fallback: use first doc name
    if md_files:
        name = os.path.splitext(md_files[0].get("name", "test_cases"))[0]
        name = re.sub(r'^[\[【][^\]】]*[\]】]', '', name).strip()
        return name or "test_cases"
    return "test_cases"


def handle_review_module_structure(args):
    """Review and suggest optimizations for test case module structure."""
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    modules = testcase_store["modules"]
    if not modules:
        return {"content": [{"type": "text", "text": "没有测试用例可供审查。请先生成用例。"}]}

    issues = []
    suggestions = []
    stats = []

    # 1. Check for empty modules/sub_modules
    for m in modules:
        subs = m.get("sub_modules", [])
        if not subs:
            issues.append(f"⚠️ 模块 '{m['name']}' 没有子模块")
        for s in subs:
            cases = s.get("test_cases", [])
            if not cases:
                issues.append(f"⚠️ 子模块 '{m['name']} > {s['name']}' 没有用例")

    # 2. Check module size balance
    module_sizes = []
    for m in modules:
        total = sum(len(s.get("test_cases", [])) for s in m.get("sub_modules", []))
        module_sizes.append((m["name"], total))
        stats.append(f"  📦 {m['name']}: {len(m.get('sub_modules', []))} 子模块, {total} 用例")

    if module_sizes:
        max_name, max_size = max(module_sizes, key=lambda x: x[1])
        min_name, min_size = min(module_sizes, key=lambda x: x[1])

        if max_size > 0 and min_size > 0 and max_size / max(min_size, 1) > 5:
            suggestions.append(
                f"💡 模块大小不均衡: '{max_name}'({max_size}个用例) vs '{min_name}'({min_size}个用例)，"
                f"建议拆分大模块或合并小模块"
            )

        # Check for overly large sub_modules (>15 cases)
        for m in modules:
            for s in m.get("sub_modules", []):
                case_count = len(s.get("test_cases", []))
                if case_count > 15:
                    suggestions.append(
                        f"💡 子模块 '{m['name']} > {s['name']}' 有 {case_count} 个用例，"
                        f"建议按场景拆分为更细粒度的子模块"
                    )

    # 3. Check for duplicate or very similar module/sub_module names
    mod_names = [m["name"] for m in modules]
    seen_names = {}
    for name in mod_names:
        key = name.strip().lower()
        if key in seen_names:
            issues.append(f"⚠️ 存在重复模块名: '{name}' 和 '{seen_names[key]}'")
        seen_names[key] = name

    for m in modules:
        sub_names = [s["name"] for s in m.get("sub_modules", [])]
        seen_sub = {}
        for name in sub_names:
            key = name.strip().lower()
            if key in seen_sub:
                issues.append(f"⚠️ 模块 '{m['name']}' 下存在重复子模块名: '{name}'")
            seen_sub[key] = name

    # 4. Check for sub_modules with only 1 case (might be too granular)
    for m in modules:
        single_case_subs = [s["name"] for s in m.get("sub_modules", [])
                           if len(s.get("test_cases", [])) == 1]
        if len(single_case_subs) >= 3:
            suggestions.append(
                f"💡 模块 '{m['name']}' 下有 {len(single_case_subs)} 个只含1个用例的子模块，"
                f"考虑合并相关子模块: {', '.join(single_case_subs[:5])}"
            )

    # 5. Check test case quality
    missing_preconditions = 0
    missing_expected = 0
    empty_steps = 0
    for m in modules:
        for s in m.get("sub_modules", []):
            for c in s.get("test_cases", []):
                if not c.get("preconditions", "").strip():
                    missing_preconditions += 1
                if not c.get("expected_result", "").strip():
                    missing_expected += 1
                if not c.get("steps") or all(not step.strip() for step in c.get("steps", [])):
                    empty_steps += 1

    if missing_preconditions > 0:
        issues.append(f"⚠️ {missing_preconditions} 个用例缺少前置条件")
    if missing_expected > 0:
        issues.append(f"⚠️ {missing_expected} 个用例缺少预期结果")
    if empty_steps > 0:
        issues.append(f"⚠️ {empty_steps} 个用例缺少执行步骤")

    # Build report
    total_cases = sum(s[1] for s in module_sizes)
    lines = [
        "📊 模块结构审查报告",
        "",
        f"总计: {len(modules)} 个模块, {total_cases} 个用例",
        "",
        "模块统计:",
    ]
    lines.extend(stats)

    if issues:
        lines.append(f"\n发现 {len(issues)} 个问题:")
        lines.extend(issues)

    if suggestions:
        lines.append(f"\n优化建议:")
        lines.extend(suggestions)

    if not issues and not suggestions:
        lines.append("\n✅ 模块结构合理，未发现明显问题。")

    lines.append(f"\n如需调整模块结构，请对需要修改的模块逐个调用 save_testcases(append_module=调整后的单个模块对象) 保存。")

    return {
        "content": [{"type": "text", "text": "\n".join(lines)}],
        "module_count": len(modules),
        "total_cases": total_cases,
        "issue_count": len(issues),
        "suggestion_count": len(suggestions),
    }


# Coverage dimension keywords for test case classification
_COVERAGE_DIMS = {
    "正向": ["正常", "正向", "成功", "默认"],
    "边界": ["边界", "最大", "最小", "上限", "下限", "空"],
    "异常": ["异常", "失败", "错误", "不存在", "无效", "非法", "超", "缺少"],
    "安全": ["安全", "认证", "授权", "权限", "注入", "xss"],
    "性能": ["并发", "性能", "大数据", "批量"],
}


def _classify_case_dimensions(case):
    """Classify a test case into coverage dimensions based on keywords."""
    title = case.get("title", "").lower()
    steps_text = " ".join(case.get("steps", [])).lower()
    combined = title + " " + steps_text
    dims = set()
    for dim, keywords in _COVERAGE_DIMS.items():
        if any(kw in combined for kw in keywords):
            dims.add(dim)
    return dims


def handle_export_report(args):
    """Generate test case report as markdown file."""
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    modules = testcase_store["modules"]
    if not modules:
        return {"content": [{"type": "text", "text": "没有测试用例可供生成报告。"}]}

    req_name = args.get("requirement_name") or _get_requirement_name()
    output_dir = args.get("output_dir", _workspace())
    # questions/confirmations from the agent about the requirements
    questions = args.get("questions", [])

    total_cases = sum(len(s.get("test_cases", []))
                      for m in modules for s in m.get("sub_modules", []))
    total_subs = sum(len(m.get("sub_modules", [])) for m in modules)

    lines = [
        "# 测试用例生成报告",
        "",
        "## 基本信息",
        "",
        "| 项目 | 内容 |",
        "|------|------|",
        f"| 需求名称 | {req_name} |",
        f"| 模块数量 | {len(modules)} |",
        f"| 子模块数量 | {total_subs} |",
        f"| 用例总数 | {total_cases} |",
        f"| XMind文件 | {req_name}_testCase.xmind |",
        "",
        "## 用例覆盖概览",
        "",
    ]

    # Per-module breakdown
    for m in modules:
        subs = m.get("sub_modules", [])
        mod_total = sum(len(s.get("test_cases", [])) for s in subs)
        lines.append(f"### {m['name']} ({mod_total} 个用例)")
        lines.append("")
        lines.append("| 子模块 | 用例数 | 覆盖维度 |")
        lines.append("|--------|--------|----------|")
        for s in subs:
            cases = s.get("test_cases", [])
            # Analyze coverage dimensions
            dimensions = set()
            for c in cases:
                dimensions.update(_classify_case_dimensions(c))
            dim_str = ", ".join(sorted(dimensions)) if dimensions else "正向"
            lines.append(f"| {s['name']} | {len(cases)} | {dim_str} |")
        lines.append("")

    # Coverage dimension summary
    all_dims = {dim: 0 for dim in _COVERAGE_DIMS}
    for m in modules:
        for s in m.get("sub_modules", []):
            for c in s.get("test_cases", []):
                for dim in _classify_case_dimensions(c):
                    all_dims[dim] += 1

    lines.append("## 覆盖维度统计")
    lines.append("")
    lines.append("| 维度 | 用例数 | 占比 |")
    lines.append("|------|--------|------|")
    for dim, count in sorted(all_dims.items(), key=lambda x: -x[1]):
        pct = f"{count / total_cases * 100:.1f}%" if total_cases > 0 else "0%"
        lines.append(f"| {dim} | {count} | {pct} |")
    lines.append("")

    # Questions / confirmation points
    if questions:
        lines.append("## 需求疑问点与确认项")
        lines.append("")
        for i, q in enumerate(questions, 1):
            lines.append(f"{i}. {q}")
        lines.append("")

    lines.append("---")
    lines.append("*报告由 TestCase Generator 自动生成*")

    report_content = "\n".join(lines)
    report_filename = f"{req_name}_testCaseReport.md"
    report_path = os.path.join(output_dir, report_filename)

    try:
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        return {
            "content": [{"type": "text", "text": f"✓ 报告已生成: {report_path}\n\n{report_content}"}],
            "report_path": report_path,
            "report_filename": report_filename,
        }
    except Exception as e:
        return {"content": [{"type": "text", "text": f"报告生成失败: {e}\n\n报告内容:\n{report_content}"}]}


def handle_export_xmind(args):
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    if not testcase_store["modules"]:
        return {"content": [{"type": "text", "text": "No test cases to export."}]}

    # Support custom naming: 需求名_testCase.xmind
    req_name = args.get("requirement_name") or _get_requirement_name()
    default_filename = f"{req_name}_testCase.xmind"
    p = args.get("output_path", os.path.join(_workspace(), default_filename))

    try:
        create_xmind_file(testcase_store["modules"], p)
        total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"]
                    for s in m.get("sub_modules", []))
        _save_phase_state("export", "completed")
        return {"content": [{"type": "text", "text": f"Exported: {p}\n{len(testcase_store['modules'])} modules, {total} cases"}],
                "xmind_path": p, "requirement_name": req_name}
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Export failed: {e}"}]}

# ============================================================
# ============================================================
# External LLM Image Processing Handlers
# ============================================================

def _get_gui_module_path():
    """Get the path to gui_llm_config.py relative to this file."""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "gui_llm_config.py")


def handle_configure_llm_api(args):
    """Launch GUI for configuring external LLM API. Runs as subprocess to avoid blocking MCP."""
    workspace = _workspace()
    gui_path = _get_gui_module_path()

    if not os.path.exists(gui_path):
        return {"content": [{"type": "text", "text": f"错误: GUI 模块未找到: {gui_path}"}]}

    try:
        # Run GUI as subprocess so it doesn't block the MCP server
        result = subprocess.run(
            [sys.executable, gui_path, workspace],
            capture_output=True, text=True, timeout=300,
            encoding='utf-8', errors='replace'
        )
        if result.returncode == 0:
            # Parse the config from stdout
            try:
                config = json.loads(result.stdout.strip())
                msg_lines = [
                    "✓ LLM API 配置已保存:",
                    f"  API 地址: {config.get('api_url', '')}",
                    f"  API Key: {'已设置' if config.get('api_key') else '未设置'}",
                    f"  模型: {config.get('model', '')}",
                    f"  多线程: {'启用 (' + str(config.get('max_threads', 3)) + ' 线程)' if config.get('enable_multithreading') else '禁用'}",
                    "",
                    "配置完成，可以调用 process_images_with_llm 开始处理图片。"
                ]
                return {"content": [{"type": "text", "text": "\n".join(msg_lines)}], "config": config}
            except json.JSONDecodeError:
                return {"content": [{"type": "text", "text": f"GUI 输出解析失败: {result.stdout[:500]}"}]}
        else:
            if "Cancelled" in (result.stdout or ""):
                return {"content": [{"type": "text", "text": "用户取消了配置。"}]}
            return {"content": [{"type": "text", "text": f"GUI 配置失败 (exit {result.returncode}): {result.stderr[:500]}"}]}
    except subprocess.TimeoutExpired:
        return {"content": [{"type": "text", "text": "配置窗口超时（5分钟），请重新调用 configure_llm_api。"}]}
    except Exception as e:
        return {"content": [{"type": "text", "text": f"启动配置窗口失败: {e}"}]}


def _process_single_image(img_info, api_url, api_key, model, prompt):
    """Process a single image with external LLM.
    
    Returns (img_id, status, result_text) where status is:
    - 'ok': successfully analyzed
    - 'skipped': LLM reported image unreadable/unclear
    - 'error': processing failed
    """
    img_id = img_info["id"]
    img_file_path = img_info["path"]
    rel_path = img_info.get("rel_path", os.path.join(".tmp", "picture", img_info["filename"]))

    # Resolve path
    if not os.path.exists(img_file_path):
        img_file_path = os.path.join(_workspace(), rel_path)
    if not os.path.exists(img_file_path):
        img_file_path = os.path.join(TMP_PIC_DIR, img_info["filename"])

    if not os.path.exists(img_file_path):
        return img_id, 'error', f"图片文件不存在: {img_file_path}"

    try:
        with open(img_file_path, 'rb') as f:
            img_data = f.read()
        mime = img_info.get("mime", "image/png")

        # Resize for LLM: max longest edge 3840, quality 90
        img_data, mime = _resize_image_for_llm(img_data, mime)

        b64 = base64.b64encode(img_data).decode('ascii')

        # Import from gui module
        gui_dir = os.path.dirname(os.path.abspath(__file__))
        if gui_dir not in sys.path:
            sys.path.insert(0, gui_dir)
        from gui_llm_config import call_llm_vision

        ok, result = call_llm_vision(api_url, api_key, model, b64, mime, prompt, timeout=120)
        if not ok:
            return img_id, 'error', result

        # Check if LLM reported the image as unreadable
        result_stripped = result.strip()
        if result_stripped.startswith(_UNREADABLE_MARKER):
            reason = result_stripped[len(_UNREADABLE_MARKER):].strip()
            return img_id, 'skipped', reason or "图片清晰度不足，无法准确解析"

        return img_id, 'ok', result
    except Exception as e:
        return img_id, 'error', f"处理异常: {e}"


# Lock for thread-safe markdown file writes (used by multi-threaded LLM processing)
_md_write_lock = threading.Lock()


def _write_image_result_to_md(img_info, analysis, skipped=False):
    """Write image analysis result back to markdown file. Thread-safe.
    
    If skipped=True, writes a skip annotation instead of analysis content.
    """
    source_doc = img_info["source_doc"]
    img_id = img_info["id"]
    md_path = os.path.join(TMP_DOC_DIR, f"{source_doc}.md")

    if not os.path.exists(md_path):
        return False, f"Markdown 文件不存在: {md_path}"

    placeholder = f"{{{{IMG:{img_id}}}}}"
    if skipped:
        replacement = f"<!-- 已跳过(图片不清晰): {img_info['filename']} — {analysis} -->"
    else:
        replacement = f"<!-- 图片分析: {img_info['filename']} -->\n{analysis}\n<!-- /图片分析 -->"

    try:
        with _md_write_lock:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()

            if placeholder in md_content:
                md_content = md_content.replace(placeholder, replacement)
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                return True, "OK"

            # Try alternate extensions
            base_id = os.path.splitext(img_id)[0]
            for ext_try in [".png", ".jpg", ".jpeg"]:
                alt_placeholder = f"{{{{IMG:{base_id}{ext_try}}}}}"
                if alt_placeholder in md_content:
                    md_content = md_content.replace(alt_placeholder, replacement)
                    with open(md_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
                    return True, "OK"

        return True, "占位符未找到，已标记为已处理"
    except Exception as e:
        return False, f"写入失败: {e}"



def handle_process_images_with_llm(args):
    """Process all pending images using configured external LLM API."""
    force_reprocess = args.get("force_reprocess", False)

    # Load LLM config
    gui_dir = os.path.dirname(os.path.abspath(__file__))
    if gui_dir not in sys.path:
        sys.path.insert(0, gui_dir)
    from gui_llm_config import load_config as load_llm_config

    config = load_llm_config(_workspace())
    api_url = config.get("api_url", "")
    api_key = config.get("api_key", "")
    model = config.get("model", "")
    enable_mt = config.get("enable_multithreading", False)
    max_threads = config.get("max_threads", 3)

    if not api_url or not model:
        return {"content": [{"type": "text", "text": (
            "未配置外部 LLM API。请先调用 configure_llm_api 打开配置窗口。"
        )}]}

    # Get pending images
    pending = testcase_store.get("pending_images", [])
    if not pending:
        _restore_store_from_cache()
        pending = testcase_store.get("pending_images", [])
    if not pending:
        return {"content": [{"type": "text", "text": "没有待处理的图片。请先调用 parse_documents 解析文档。"}]}

    # Filter to unprocessed images (or all if force_reprocess)
    to_process = [img for img in pending if not img["processed"] or force_reprocess]
    if not to_process:
        return {"content": [{"type": "text", "text": "所有图片已处理完毕。如需重新处理，请传入 force_reprocess=true。"}]}

    prompt = IMAGE_ANALYSIS_PROMPT
    results_log = []
    success_count = 0
    fail_count = 0
    skip_count = 0

    if enable_mt and len(to_process) > 1:
        # Multi-threaded processing
        import concurrent.futures
        actual_threads = min(max_threads, len(to_process))
        results_log.append(f"🚀 多线程模式: {actual_threads} 线程并发处理 {len(to_process)} 张图片\n")

        with concurrent.futures.ThreadPoolExecutor(max_workers=actual_threads) as executor:
            future_to_img = {}
            for img in to_process:
                future = executor.submit(_process_single_image, img, api_url, api_key, model, prompt)
                future_to_img[future] = img

            for future in concurrent.futures.as_completed(future_to_img):
                img = future_to_img[future]
                try:
                    img_id, status, result_text = future.result()
                    if status == 'ok':
                        wrote_ok, write_msg = _write_image_result_to_md(img, result_text)
                        img["processed"] = True
                        success_count += 1
                        if not wrote_ok:
                            results_log.append(f"  ⚠ [{img_id}] LLM成功但写入失败: {write_msg}")
                        else:
                            results_log.append(f"  ✓ [{img_id}] 处理成功")
                    elif status == 'skipped':
                        _write_image_result_to_md(img, result_text, skipped=True)
                        img["processed"] = True
                        img["skipped"] = True
                        skip_count += 1
                        results_log.append(f"  ⏭ [{img_id}] 已跳过: {result_text[:100]}")
                    else:
                        fail_count += 1
                        results_log.append(f"  ✗ [{img_id}] {result_text[:200]}")
                except Exception as e:
                    fail_count += 1
                    results_log.append(f"  ✗ [{img['id']}] 异常: {e}")
    else:
        # Sequential processing
        results_log.append(f"📝 顺序处理 {len(to_process)} 张图片\n")
        for i, img in enumerate(to_process, 1):
            results_log.append(f"  [{i}/{len(to_process)}] 处理 {img['id']} ...")
            img_id, status, result_text = _process_single_image(img, api_url, api_key, model, prompt)
            if status == 'ok':
                wrote_ok, write_msg = _write_image_result_to_md(img, result_text)
                img["processed"] = True
                success_count += 1
                if not wrote_ok:
                    results_log.append(f"    ⚠ LLM成功但写入失败: {write_msg}")
                else:
                    results_log.append("    ✓ 成功")
            elif status == 'skipped':
                _write_image_result_to_md(img, result_text, skipped=True)
                img["processed"] = True
                img["skipped"] = True
                skip_count += 1
                results_log.append(f"    ⏭ 已跳过: {result_text[:100]}")
            else:
                fail_count += 1
                results_log.append(f"    ✗ {result_text[:200]}")

    # Persist progress
    _sync_store_to_cache()
    total = len(pending)
    processed = sum(1 for p in pending if p["processed"])
    skipped_total = sum(1 for p in pending if p.get("skipped"))
    _save_phase_state("image_analysis", "completed" if processed == total else "in_progress", {
        "total": total,
        "processed": processed,
        "skipped": skipped_total,
    })

    results_log.append(f"\n处理完成: {success_count} 成功, {skip_count} 跳过(不清晰), {fail_count} 失败, 总进度 {processed}/{total}")
    if skipped_total > 0:
        results_log.append(f"  ⏭ 共 {skipped_total} 张图片因清晰度问题被跳过，不影响后续用例生成。")
    if processed == total:
        results_log.append("\n所有图片已处理完毕！请调用 get_doc_summary 获取文档结构概览。")
    elif fail_count > 0:
        results_log.append("\n部分图片处理失败，可以重新调用 process_images_with_llm 重试。")

    return {
        "content": [{"type": "text", "text": "\n".join(results_log)}],
        "success_count": success_count,
        "skip_count": skip_count,
        "fail_count": fail_count,
        "total_processed": processed,
        "total_images": total,
    }


# ============================================================
# MCP Main Loop
# ============================================================

HANDLERS = {
    "setup_environment": handle_setup_environment,
    "clear_cache": handle_clear_cache,
    "parse_documents": handle_parse_documents,
    "get_workflow_state": handle_get_workflow_state,
    "get_doc_summary": handle_get_doc_summary,
    "get_doc_section": handle_get_doc_section,
    "get_parsed_markdown": handle_get_parsed_markdown,
    "save_testcases": handle_save_testcases,
    "get_testcases": handle_get_testcases,
    "export_xmind": handle_export_xmind,
    "review_module_structure": handle_review_module_structure,
    "export_report": handle_export_report,
    "configure_llm_api": handle_configure_llm_api,
    "process_images_with_llm": handle_process_images_with_llm,
}


def handle_request(req):
    method = req.get("method", "")
    rid = req.get("id")
    params = req.get("params", {})

    sys.stderr.write(f"[MCP] Received: method={method} id={rid}\n")
    sys.stderr.flush()

    if method == "initialize":
        send_response(rid, {
            "protocolVersion": "2024-11-05",
            "capabilities": {"tools": {"listChanged": False}},
            "serverInfo": {"name": "testcase-generator", "version": "7.0.0"}
        })
    elif method == "notifications/initialized":
        pass  # No response needed for notifications
    elif method == "notifications/cancelled":
        pass  # No response needed for cancellation notifications
    elif method == "tools/list":
        send_response(rid, {"tools": TOOLS})
    elif method == "tools/call":
        name = params.get("name", "")
        handler = HANDLERS.get(name)
        if handler:
            try:
                result = handler(params.get("arguments", {}))
                send_response(rid, result)
            except Exception as e:
                tb = traceback.format_exc()
                sys.stderr.write(f"[MCP] Tool error in {name}: {tb}\n")
                sys.stderr.flush()
                send_response(rid, {"content": [{"type": "text", "text": f"Error in {name}: {e}"}], "isError": True})
        else:
            send_error(rid, -32601, f"Unknown tool: {name}")
    elif method == "ping":
        send_response(rid, {})
    else:
        if rid is not None:
            send_error(rid, -32601, f"Method not found: {method}")


def main():
    sys.stderr.write("TestCase Generator MCP Server v7.0 starting...\n")
    sys.stderr.write(f"  Workspace: {_INITIAL_WORKSPACE}\n")
    sys.stderr.write(f"  Python: {sys.version.split()[0]}\n")
    sys.stderr.flush()

    while True:
        try:
            line = sys.stdin.buffer.readline()
            if not line:
                sys.stderr.write("[MCP] stdin closed, shutting down.\n")
                sys.stderr.flush()
                return

            stripped = line.strip()
            if not stripped:
                continue

            decoded = stripped.decode('utf-8', errors='replace')

            if decoded.startswith('{'):
                try:
                    request = json.loads(decoded)
                    handle_request(request)
                except json.JSONDecodeError as e:
                    sys.stderr.write(f"[MCP] JSON parse error: {e}\n")
                    sys.stderr.flush()
            elif decoded.lower().startswith('content-length'):
                try:
                    cl = int(decoded.split(":", 1)[1].strip())
                except (ValueError, IndexError):
                    continue
                sys.stdin.buffer.readline()  # skip blank line
                body = b""
                while len(body) < cl:
                    chunk = sys.stdin.buffer.read(cl - len(body))
                    if not chunk:
                        break
                    body += chunk
                if body:
                    try:
                        request = json.loads(body.decode('utf-8'))
                        handle_request(request)
                    except json.JSONDecodeError as e:
                        sys.stderr.write(f"[MCP] JSON parse error: {e}\n")
                        sys.stderr.flush()
        except KeyboardInterrupt:
            sys.stderr.write("[MCP] Interrupted, shutting down.\n")
            sys.stderr.flush()
            return
        except Exception as e:
            sys.stderr.write(f"[MCP] Error in main loop: {e}\n")
            sys.stderr.flush()


if __name__ == "__main__":
    main()
