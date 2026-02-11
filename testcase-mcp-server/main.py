#!/usr/bin/env python3
"""
Test Case Generator MCP Server v6.0
- Phase-based workflow with file cache for cross-session resume
- Splits document reading into summary + sections to reduce context
- Session switch hints when image batch threshold reached
- All state persisted to .tmp/cache/ for recovery
"""

import sys
import json
import os
import glob
import zipfile
import re
import subprocess
import base64
import hashlib
import shutil
from io import BytesIO
from xml.etree import ElementTree as ET

# ============================================================
# Windows binary mode for stdin/stdout
# ============================================================
if sys.platform == "win32":
    import msvcrt
    msvcrt.setmode(sys.stdin.fileno(), os.O_BINARY)
    msvcrt.setmode(sys.stdout.fileno(), os.O_BINARY)

sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# ============================================================
# Auto-install helper
# ============================================================

def _ensure_pkg(import_name, pip_name):
    try:
        __import__(import_name)
        return True
    except ImportError:
        try:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", pip_name, "-q"],
                stdout=subprocess.DEVNULL, stderr=subprocess.PIPE)
            return True
        except Exception:
            return False

# ============================================================
# MCP Protocol (stdio JSON-RPC)
# ============================================================

def send_response(rid, result):
    msg = json.dumps({"jsonrpc": "2.0", "id": rid, "result": result}, ensure_ascii=False)
    raw = msg.encode('utf-8')
    sys.stdout.buffer.write(raw + b"\n")
    sys.stdout.buffer.flush()


def send_error(rid, code, message):
    msg = json.dumps({"jsonrpc": "2.0", "id": rid, "error": {"code": code, "message": message}}, ensure_ascii=False)
    raw = msg.encode('utf-8')
    sys.stdout.buffer.write(raw + b"\n")
    sys.stdout.buffer.flush()


# ============================================================
# Constants & Paths
# ============================================================

WORKSPACE_DIR = None
TMP_DOC_DIR = os.path.join(os.getcwd(), ".tmp", "doc_mk")
TMP_PIC_DIR = os.path.join(os.getcwd(), ".tmp", "picture")
TMP_CACHE_DIR = os.path.join(os.getcwd(), ".tmp", "cache")

# No forced session switch — let the system decide naturally.
# Cross-session resume is still fully supported via .tmp/cache/.


def _update_workspace(directory):
    global WORKSPACE_DIR, TMP_DOC_DIR, TMP_PIC_DIR, TMP_CACHE_DIR
    WORKSPACE_DIR = directory
    TMP_DOC_DIR = os.path.join(directory, ".tmp", "doc_mk")
    TMP_PIC_DIR = os.path.join(directory, ".tmp", "picture")
    TMP_CACHE_DIR = os.path.join(directory, ".tmp", "cache")


def _workspace():
    return WORKSPACE_DIR or os.getcwd()

# ============================================================
# Cache / Persistence Layer
# ============================================================

def _cache_path(filename):
    return os.path.join(TMP_CACHE_DIR, filename)


def _save_cache(filename, data):
    os.makedirs(TMP_CACHE_DIR, exist_ok=True)
    path = _cache_path(filename)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_cache(filename, default=None):
    path = _cache_path(filename)
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass
    return default if default is not None else {}


def _save_phase_state(phase, status, extra=None):
    """Update phase_state.json with current progress."""
    state = _load_cache("phase_state.json", {
        "current_phase": phase,
        "workspace_dir": _workspace(),
        "phases": {}
    })
    state["current_phase"] = phase
    state["workspace_dir"] = _workspace()
    state.setdefault("phases", {})
    if phase not in state["phases"]:
        state["phases"][phase] = {}
    state["phases"][phase]["status"] = status
    if extra:
        state["phases"][phase].update(extra)
    _save_cache("phase_state.json", state)
    return state


def _reset_phase_state():
    """Clear all phase state for a fresh start."""
    state = {
        "current_phase": "init",
        "workspace_dir": _workspace(),
        "phases": {}
    }
    _save_cache("phase_state.json", state)
    return state

# ============================================================
# Image helpers
# ============================================================

def _img_mime(ext):
    return {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff",
            "emf": "image/emf", "wmf": "image/wmf"}.get(ext.lower().lstrip('.'), "image/png")


def _generate_image_id(doc_name, img_name):
    raw = f"{doc_name}_{img_name}"
    short_hash = hashlib.md5(raw.encode()).hexdigest()[:8]
    base = os.path.splitext(img_name)[0]
    ext = os.path.splitext(img_name)[1]
    return f"{base}_{short_hash}{ext}"


def _resize_image(img_data, ext):
    final_data = img_data
    final_mime = _img_mime(ext)
    try:
        from PIL import Image
        img_obj = Image.open(BytesIO(img_data))
        w, h = img_obj.size

        # Skip tiny images (icons, decorations) - no analysis value
        if w < 60 or h < 60:
            return final_data, final_mime

        # Resize large images, scale based on content density
        MAX_DIM = 1568  # Claude vision optimal tile boundary (multiple of 784)
        if w > MAX_DIM or h > MAX_DIM:
            ratio = min(MAX_DIM / w, MAX_DIM / h)
            new_w, new_h = int(w * ratio), int(h * ratio)
            img_obj = img_obj.resize((new_w, new_h), Image.LANCZOS)

        # Force grayscale for document images (tables, flowcharts, text)
        # Requirement docs rarely need color; grayscale saves ~60% base64 size
        if img_obj.mode not in ('L', 'LA'):
            img_obj = img_obj.convert('L')

        buf = BytesIO()
        if img_obj.mode in ('RGBA', 'P', 'LA'):
            img_obj.save(buf, format='PNG', optimize=True)
            final_mime = "image/png"
        else:
            # quality=65 is the sweet spot: text stays sharp, ~30% smaller than q80
            img_obj.save(buf, format='JPEG', quality=65)
            final_mime = "image/jpeg"
        final_data = buf.getvalue()
    except Exception:
        pass
    return final_data, final_mime

# ============================================================
# DOCX → Markdown + Images extraction
# ============================================================

def _build_rid_to_media(filepath):
    rid_to_media = {}
    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            rels_path = 'word/_rels/document.xml.rels'
            if rels_path in z.namelist():
                rels_root = ET.fromstring(z.read(rels_path))
                for rel in rels_root:
                    target = rel.get('Target', '')
                    rid = rel.get('Id', '')
                    if 'media/' in target:
                        rid_to_media[rid] = target.split('/')[-1]
    except Exception:
        pass
    return rid_to_media


def _find_images_in_element(elem, rid_to_media):
    A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    found = []
    seen = set()
    for blip in elem.iter(f'{{{A_NS}}}blip'):
        rid = blip.get(f'{{{R_NS}}}embed', '')
        if rid and rid in rid_to_media:
            name = rid_to_media[rid]
            if name not in seen:
                seen.add(name)
                found.append(name)
    for pict in elem.iter(f'{{{W_NS}}}pict'):
        for child in pict.iter():
            rid = child.get(f'{{{R_NS}}}id', '')
            if rid and rid in rid_to_media:
                name = rid_to_media[rid]
                if name not in seen:
                    seen.add(name)
                    found.append(name)
    return found


def _detect_heading_level(para):
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    if para.style and para.style.name and para.style.name.startswith('Heading'):
        lvl_str = para.style.name.replace('Heading ', '').replace('Heading', '1')
        try:
            return int(lvl_str)
        except ValueError:
            return 1
    try:
        pPr = para._element.find(f'{W_NS}pPr')
        if pPr is not None:
            outlineLvl = pPr.find(f'{W_NS}outlineLvl')
            if outlineLvl is not None:
                val = int(outlineLvl.get(f'{W_NS}val', '-1'))
                if val >= 0:
                    return val + 1
    except Exception:
        pass
    max_size = 0
    is_bold = False
    for run in para.runs:
        if run.bold:
            is_bold = True
        if run.font.size:
            sz = run.font.size / 12700
            if sz > max_size:
                max_size = sz
    if max_size >= 22:
        return 1
    if max_size >= 17 and is_bold:
        return 1
    if max_size >= 15 and is_bold:
        return 2
    if max_size >= 14 and is_bold:
        return 3
    return 0

def _table_to_markdown(table, rid_to_media, doc_name, image_registry):
    rows = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
            cell_images = []
            for para in cell.paragraphs:
                imgs = _find_images_in_element(para._element, rid_to_media)
                cell_images.extend(imgs)
            if cell_images:
                for img_name in cell_images:
                    img_id = _generate_image_id(doc_name, img_name)
                    image_registry[img_name] = img_id
                    cell_text += f" {{{{IMG:{img_id}}}}}"
            row_data.append(cell_text)
        rows.append(row_data)
    if not rows:
        return ""

    # Detect single-column tables containing code/diagrams
    if len(rows[0]) == 1:
        content = rows[0][0].replace('\\|', '|')
        # Check for Mermaid diagrams, SQL, code blocks
        code_indicators = [
            'sequenceDiagram', 'stateDiagram', 'erDiagram', 'flowchart',
            'classDiagram', 'gantt', 'pie', 'graph ',
            'SELECT ', 'INSERT ', 'CREATE TABLE', 'ALTER TABLE', 'DROP TABLE',
            'DELETE FROM', 'UPDATE ',
        ]
        lang_hints = {
            'sequenceDiagram': 'mermaid', 'stateDiagram': 'mermaid',
            'erDiagram': 'mermaid', 'flowchart': 'mermaid',
            'classDiagram': 'mermaid', 'gantt': 'mermaid', 'pie': 'mermaid',
            'graph ': 'mermaid',
            'SELECT ': 'sql', 'INSERT ': 'sql', 'CREATE TABLE': 'sql',
            'ALTER TABLE': 'sql', 'DROP TABLE': 'sql', 'DELETE FROM': 'sql',
            'UPDATE ': 'sql',
        }
        for indicator in code_indicators:
            if indicator in content:
                lang = lang_hints.get(indicator, '')
                # Restore newlines for readability
                raw_text = table.rows[0].cells[0].text.strip()
                return f"```{lang}\n{raw_text}\n```"

        # Check for Java/JSON/XML/Plaintext code patterns
        code_patterns = [
            ('Java ', 'java'), ('JSON ', 'json'), ('XML ', 'xml'),
            ('Plaintext ', 'text'), ('String ', 'java'),
            ('public ', 'java'), ('private ', 'java'),
            ('if (', 'java'), ('if(', 'java'),
        ]
        for pattern, lang in code_patterns:
            if content.startswith(pattern) or content.startswith(pattern.lower()):
                raw_text = table.rows[0].cells[0].text.strip()
                return f"```{lang}\n{raw_text}\n```"

    md_lines = []
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        while len(row) < len(rows[0]):
            row.append("")
        md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
    return "\n".join(md_lines)


def convert_docx_to_markdown(filepath):
    doc_name = os.path.splitext(os.path.basename(filepath))[0]
    rid_to_media = _build_rid_to_media(filepath)
    image_registry = {}
    image_data_map = {}
    md_lines = [f"# {doc_name}", ""]

    try:
        from docx import Document
        doc = Document(filepath)
    except ImportError:
        return _convert_docx_raw(filepath)
    except Exception:
        return _convert_docx_raw(filepath)

    # Iterate body elements in order to preserve table positions
    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue
            text = para.text.strip()
            para_images = _find_images_in_element(para._element, rid_to_media)
            if text:
                level = _detect_heading_level(para)
                if level > 0:
                    md_lines.append("")
                    md_lines.append(f"{'#' * (level + 1)} {text}")
                    md_lines.append("")
                else:
                    md_lines.append(text)
                    md_lines.append("")
            if para_images:
                for img_name in para_images:
                    img_id = _generate_image_id(doc_name, img_name)
                    image_registry[img_name] = img_id
                    md_lines.append(f"{{{{IMG:{img_id}}}}}")
                    md_lines.append("")
        elif tag == 'tbl':
            # Find the corresponding table object
            tbl = None
            for t in doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl is not None:
                md_lines.append("")
                table_md = _table_to_markdown(tbl, rid_to_media, doc_name, image_registry)
                if table_md:
                    md_lines.append(table_md)
                    md_lines.append("")

    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            skipped_ids = []
            for img_path in media:
                name = os.path.basename(img_path)
                if name in image_registry:
                    img_id = image_registry[name]
                    ext = os.path.splitext(name)[1].lstrip('.')
                    img_data = z.read(img_path)
                    if len(img_data) >= 500 and ext.lower() not in ('emf', 'wmf'):
                        # Skip tiny images by pixel dimensions (icons/decorations)
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            _w, _h = _tmp.size
                            if _w < 60 or _h < 60:
                                skipped_ids.append(img_id)
                                continue
                        except Exception:
                            pass
                        image_data_map[img_id] = (img_data, ext)
                    else:
                        skipped_ids.append(img_id)
            # Replace skipped image placeholders with annotation in markdown
            if skipped_ids:
                md_text = "\n".join(md_lines)
                for sid in skipped_ids:
                    placeholder = f"{{{{IMG:{sid}}}}}"
                    annotation = f"<!-- 已跳过小图标: {sid} -->"
                    md_text = md_text.replace(placeholder, annotation)
                md_lines = md_text.split("\n")
    except Exception:
        pass

    return "\n".join(md_lines), image_registry, image_data_map

def _convert_docx_raw(filepath):
    doc_name = os.path.splitext(os.path.basename(filepath))[0]
    rid_to_media = _build_rid_to_media(filepath)
    image_registry = {}
    image_data_map = {}
    md_lines = [f"# {doc_name}", ""]

    try:
        with zipfile.ZipFile(filepath, 'r') as z:
            W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            if 'word/document.xml' in z.namelist():
                root = ET.fromstring(z.read('word/document.xml'))
                body = root.find(f'{W}body')
                if body is None:
                    body = root
                for child in body:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'p':
                        texts = [t.text for t in child.iter(f'{W}t') if t.text]
                        text = ''.join(texts).strip()
                        is_heading = False
                        if text:
                            pPr = child.find(f'{W}pPr')
                            if pPr is not None:
                                pStyle = pPr.find(f'{W}pStyle')
                                if pStyle is not None:
                                    sv = pStyle.get(f'{W}val', '')
                                    if 'Heading' in sv or 'heading' in sv:
                                        is_heading = True
                                        m = re.search(r'\d+', sv)
                                        level = int(m.group()) if m else 1
                                        md_lines.append("")
                                        md_lines.append(f"{'#' * (level + 1)} {text}")
                                        md_lines.append("")
                            if not is_heading:
                                md_lines.append(text)
                                md_lines.append("")
                        para_images = _find_images_in_element(child, rid_to_media)
                        for img_name in para_images:
                            img_id = _generate_image_id(doc_name, img_name)
                            image_registry[img_name] = img_id
                            md_lines.append(f"{{{{IMG:{img_id}}}}}")
                            md_lines.append("")
                    elif tag == 'tbl':
                        rows = []
                        for tr in child.iter(f'{W}tr'):
                            row = []
                            for tc in tr.iter(f'{W}tc'):
                                cell_text = ''.join(t.text for t in tc.iter(f'{W}t') if t.text).strip()
                                cell_images = _find_images_in_element(tc, rid_to_media)
                                for img_name in cell_images:
                                    img_id = _generate_image_id(doc_name, img_name)
                                    image_registry[img_name] = img_id
                                    cell_text += f" {{{{IMG:{img_id}}}}}"
                                row.append(cell_text.replace('|', '\\|'))
                            if row:
                                rows.append(row)
                        if rows:
                            md_lines.append("")
                            md_lines.append("| " + " | ".join(rows[0]) + " |")
                            md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                            for row in rows[1:]:
                                while len(row) < len(rows[0]):
                                    row.append("")
                                md_lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")
                            md_lines.append("")
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            skipped_ids = []
            for img_path in media:
                name = os.path.basename(img_path)
                if name in image_registry:
                    img_id = image_registry[name]
                    ext = os.path.splitext(name)[1].lstrip('.')
                    img_data = z.read(img_path)
                    if len(img_data) >= 500 and ext.lower() not in ('emf', 'wmf'):
                        try:
                            from PIL import Image as _Img
                            _tmp = _Img.open(BytesIO(img_data))
                            _w, _h = _tmp.size
                            if _w < 60 or _h < 60:
                                skipped_ids.append(img_id)
                                continue
                        except Exception:
                            pass
                        image_data_map[img_id] = (img_data, ext)
                    else:
                        skipped_ids.append(img_id)
            # Replace skipped image placeholders with annotation in markdown
            if skipped_ids:
                md_text = "\n".join(md_lines)
                for sid in skipped_ids:
                    placeholder = f"{{{{IMG:{sid}}}}}"
                    annotation = f"<!-- 已跳过小图标: {sid} -->"
                    md_text = md_text.replace(placeholder, annotation)
                md_lines = md_text.split("\n")
    except Exception as e:
        md_lines.append(f"\n> 解析错误: {e}\n")

    return "\n".join(md_lines), image_registry, image_data_map

# ============================================================
# In-memory Store (backed by cache files)
# ============================================================

testcase_store = {
    "modules": [],
    "pending_images": [],
    "md_files": [],
    "session_image_count": 0,  # images processed in current session (not persisted)
}


def _sync_store_to_cache():
    """Persist critical store data to cache files."""
    _save_cache("image_progress.json", {
        "pending_images": testcase_store["pending_images"],
        "md_files": testcase_store["md_files"],
    })
    # Always write testcases (even empty) to avoid stale cache
    _save_cache("testcases.json", {
        "modules": testcase_store["modules"],
    })


def _restore_store_from_cache():
    """Restore store from cache files (for session resume)."""
    img_data = _load_cache("image_progress.json")
    if img_data:
        testcase_store["pending_images"] = img_data.get("pending_images", [])
        testcase_store["md_files"] = img_data.get("md_files", [])

    tc_data = _load_cache("testcases.json")
    if tc_data:
        testcase_store["modules"] = tc_data.get("modules", [])

# ============================================================
# Document Section Parser (for get_doc_summary / get_doc_section)
# ============================================================

def _parse_md_sections(md_content):
    """Parse markdown into sections by headings. Returns list of {heading, level, start, end, char_count}."""
    lines = md_content.split('\n')
    sections = []
    current = None
    in_code_block = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        # Track code blocks to avoid treating # comments as headings
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        if stripped.startswith('#') and not stripped.startswith('#!'):
            hashes = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            if not title:  # Skip lines that are just '#' with no title
                continue
            if current:
                current["end"] = i
                current["char_count"] = sum(len(lines[j]) for j in range(current["start"], i))
            current = {"heading": title, "level": hashes, "start": i, "end": len(lines), "char_count": 0}
            sections.append(current)

    if current:
        current["end"] = len(lines)
        current["char_count"] = sum(len(lines[j]) for j in range(current["start"], current["end"]))

    # If no headings found, treat entire content as one section
    if not sections:
        sections.append({
            "heading": "(全文)",
            "level": 1,
            "start": 0,
            "end": len(lines),
            "char_count": len(md_content)
        })

    return sections


def _build_doc_summary():
    """Build summary of all markdown docs: structure tree + stats."""
    md_files = testcase_store.get("md_files", [])
    summary = {"documents": [], "total_chars": 0, "total_sections": 0}

    for md_info in md_files:
        try:
            with open(md_info["path"], 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception:
            continue

        sections = _parse_md_sections(content)
        doc_entry = {
            "name": md_info["name"],
            "total_chars": len(content),
            "sections": []
        }
        for sec in sections:
            doc_entry["sections"].append({
                "heading": sec["heading"],
                "level": sec["level"],
                "char_count": sec["char_count"],
                "line_start": sec["start"],
                "line_end": sec["end"],
            })
        summary["documents"].append(doc_entry)
        summary["total_chars"] += len(content)
        summary["total_sections"] += len(sections)

    # Also save to cache
    _save_cache("doc_summary.json", summary)
    return summary

# ============================================================
# XMind Export
# ============================================================

def _esc(text):
    if not text:
        return ""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def create_xmind_file(modules, output_path):
    counter = [0]
    def nid():
        counter[0] += 1
        return f"t{counter[0]}"
    def topic(title, children=None):
        x = f'<topic id="{nid()}"><title>{_esc(title)}</title>'
        if children:
            x += '<children><topics type="attached">' + ''.join(children) + '</topics></children>'
        return x + '</topic>'

    mod_topics = []
    for m in modules:
        sub_topics = []
        for s in m.get("sub_modules", []):
            case_topics = []
            for c in s.get("test_cases", []):
                # Build chain: 用例标题 → 前置条件 → 执行步骤 → 预期结果
                inner = None
                if c.get("expected_result"):
                    inner = topic(f"预期结果: {c['expected_result']}")
                steps = c.get("steps", [])
                if steps:
                    steps_text = "\n".join(f"{i}. {step}" for i, step in enumerate(steps, 1))
                    inner = topic(f"执行步骤:\n{steps_text}", [inner] if inner else None)
                if c.get("preconditions"):
                    inner = topic(f"前置条件: {c['preconditions']}", [inner] if inner else None)
                case_topics.append(topic(c.get("title", "未命名用例"), [inner] if inner else None))
            sub_topics.append(topic(s.get("name", ""), case_topics))
        mod_topics.append(topic(m.get("name", ""), sub_topics))

    content_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<xmap-content xmlns="urn:xmind:xmap:xmlns:content:2.0"
  xmlns:fo="http://www.w3.org/1999/XSL/Format"
  xmlns:svg="http://www.w3.org/2000/svg"
  xmlns:xhtml="http://www.w3.org/1999/xhtml"
  xmlns:xlink="http://www.w3.org/1999/xlink" version="2.0">
  <sheet id="sheet_1"><title>测试用例</title>{topic("测试用例", mod_topics)}</sheet>
</xmap-content>'''

    manifest_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="no"?>
<manifest xmlns="urn:xmind:xmap:xmlns:manifest:1.0">
  <file-entry full-path="content.xml" media-type="text/xml"/>
  <file-entry full-path="META-INF/" media-type=""/>
  <file-entry full-path="META-INF/manifest.xml" media-type="text/xml"/>
</manifest>'''

    meta_xml = ('<?xml version="1.0" encoding="UTF-8"?>'
                '<meta xmlns="urn:xmind:xmap:xmlns:meta:2.0" version="2.0">'
                '<Author><Name>TestCase Generator</Name></Author></meta>')

    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('content.xml', content_xml)
        zf.writestr('META-INF/manifest.xml', manifest_xml)
        zf.writestr('meta.xml', meta_xml)
    return output_path

# ============================================================
# MCP Tool Definitions
# ============================================================

TOOLS = [
    {
        "name": "setup_environment",
        "description": "Check and auto-install Python dependencies (python-docx, Pillow). No OCR needed.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "parse_documents",
        "description": "Parse .docx files from doc/ directory: convert to markdown (.tmp/doc_mk/), extract images to .tmp/picture/. Returns file list and pending image count. State is persisted to .tmp/cache/ for cross-session resume. Will block if there's an in-progress workflow (pass force=true to override).",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Directory containing .docx files (default: cwd)"},
                "file_patterns": {"type": "string", "description": "Glob pattern (default: *.docx)"},
                "force": {"type": "boolean", "description": "Force re-parse even if there's in-progress work (default: false)"}
            },
            "required": []
        }
    },
    {
        "name": "get_pending_image",
        "description": "Get the next unprocessed image for vision analysis. Returns one image at a time (base64). After analyzing, call submit_image_result with the result. Progress is auto-saved for cross-session resume.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "submit_image_result",
        "description": "Submit vision analysis result for an image. Writes result back to markdown file replacing {{IMG:id}} placeholder. Progress is auto-saved to cache.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "image_id": {"type": "string", "description": "The unique image ID (from get_pending_image)"},
                "analysis": {"type": "string", "description": "The analysis result text describing image content"}
            },
            "required": ["image_id", "analysis"]
        }
    },
    {
        "name": "get_workflow_state",
        "description": "Get current workflow state for session resume. Returns phase progress, pending work, and resume instructions. Call this at the start of a new session to continue previous work.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "directory": {"type": "string", "description": "Workspace directory (default: cwd). Used to locate .tmp/cache/."}
            },
            "required": []
        }
    },
    {
        "name": "get_doc_summary",
        "description": "Get document structure summary (heading tree + char counts per section) without loading full content. Use this to plan which sections to read with get_doc_section.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "get_doc_section",
        "description": "Read a specific section of a markdown document by heading name. Returns only that section's content, reducing context usage vs loading the full document.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "doc_name": {"type": "string", "description": "Markdown filename (e.g. 'xxx.md')"},
                "section_heading": {"type": "string", "description": "Heading text to match (fuzzy match supported)"},
                "include_subsections": {"type": "boolean", "description": "Include child sections (default: true)"}
            },
            "required": ["doc_name"]
        }
    },
    {
        "name": "get_parsed_markdown",
        "description": "Read all processed markdown files. WARNING: may be large. Prefer get_doc_summary + get_doc_section for large documents to avoid context overflow.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "save_testcases",
        "description": "Save test cases. Also persists to .tmp/cache/testcases.json for cross-session access. Supports incremental save via append_module parameter.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "modules": {"type": "array", "description": "Test case module list", "items": {"type": "object"}},
                "append_module": {"type": "object", "description": "Single module to append to existing cases (for incremental generation)"}
            },
            "required": []
        }
    },
    {
        "name": "get_testcases",
        "description": "Get all current test cases. Loads from cache if memory is empty.",
        "inputSchema": {"type": "object", "properties": {}, "required": []}
    },
    {
        "name": "export_xmind",
        "description": "Export test cases to XMind format.",
        "inputSchema": {
            "type": "object",
            "properties": {
                "output_path": {"type": "string", "description": "Output file path (default: test_cases.xmind)"}
            },
            "required": []
        }
    }
]

# ============================================================
# Tool Handlers
# ============================================================

def handle_setup_environment(args):
    results = []
    all_ok = True
    results.append(f"Python {sys.version.split()[0]}")
    deps = {"docx": "python-docx", "PIL": "Pillow"}
    for imp, pip_name in deps.items():
        try:
            __import__(imp)
            results.append(f"  [ok] {pip_name}")
        except ImportError:
            results.append(f"  [installing] {pip_name}...")
            if _ensure_pkg(imp, pip_name):
                results.append(f"  [ok] {pip_name} installed")
            else:
                results.append(f"  [FAIL] {pip_name}")
                all_ok = False
    results.append("")
    results.append("OK - environment ready" if all_ok else "WARN - some deps failed")
    return {"content": [{"type": "text", "text": "\n".join(results)}], "all_ok": all_ok}


def handle_parse_documents(args):
    directory = args.get("directory", os.getcwd())
    pattern = args.get("file_patterns", "*.docx")
    force = args.get("force", False)
    _update_workspace(directory)

    # Default: look in doc/ subdirectory
    doc_dir = os.path.join(directory, "doc")
    search_dir = doc_dir if os.path.isdir(doc_dir) else directory

    # Protection: check if there's an in-progress workflow
    if not force:
        existing_state = _load_cache("phase_state.json")
        if existing_state:
            phases = existing_state.get("phases", {})
            img_phase = phases.get("image_analysis", {})
            if img_phase.get("status") == "in_progress":
                processed = img_phase.get("processed", 0)
                total = img_phase.get("total", 0)
                return {"content": [{"type": "text", "text": (
                    f"⚠️ 检测到未完成的图片处理进度 ({processed}/{total})。\n"
                    f"重新解析会丢失已有进度。如需继续处理，请调用 get_workflow_state 恢复。\n"
                    f"如确认要重新开始，请传入 force=true 参数。"
                )}], "blocked": True}

    files = glob.glob(os.path.join(search_dir, "**", pattern), recursive=True)
    if not files:
        files = glob.glob(os.path.join(search_dir, pattern))
    if not files:
        return {"content": [{"type": "text", "text": f"No .docx files found in {search_dir}"}]}

    # Clean doc/picture dirs and reset all cache state for fresh start
    for d in (TMP_DOC_DIR, TMP_PIC_DIR):
        if os.path.exists(d):
            shutil.rmtree(d)
        os.makedirs(d, exist_ok=True)
    os.makedirs(TMP_CACHE_DIR, exist_ok=True)

    # Reset all cache state files when re-parsing
    for cache_file in ("phase_state.json", "image_progress.json", "testcases.json", "doc_summary.json"):
        cache_fp = _cache_path(cache_file)
        if os.path.exists(cache_fp):
            os.remove(cache_fp)

    # Reset in-memory store
    testcase_store["modules"] = []
    testcase_store["pending_images"] = []
    testcase_store["md_files"] = []
    testcase_store["session_image_count"] = 0

    # Initialize clean phase state
    _reset_phase_state()

    all_md_files = []
    all_pending_images = []
    content_parts = []

    for fpath in files:
        try:
            doc_name = os.path.splitext(os.path.basename(fpath))[0]
            md_text, image_registry, image_data_map = convert_docx_to_markdown(fpath)

            md_filename = f"{doc_name}.md"
            md_path = os.path.join(TMP_DOC_DIR, md_filename)
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            all_md_files.append({"name": md_filename, "path": md_path})

            skipped = 0
            for img_id, (img_data, ext) in image_data_map.items():
                resized_data, mime = _resize_image(img_data, ext)
                out_ext = ".png" if "png" in mime else ".jpg"
                img_filename = os.path.splitext(img_id)[0] + out_ext
                img_path = os.path.join(TMP_PIC_DIR, img_filename)
                with open(img_path, 'wb') as f:
                    f.write(resized_data)
                all_pending_images.append({
                    "id": img_id, "filename": img_filename, "path": img_path,
                    "mime": mime, "size": len(img_data), "source_doc": doc_name,
                    "processed": False
                })

            for orig_name, uid in image_registry.items():
                if uid not in image_data_map:
                    skipped += 1

            content_parts.append({
                "type": "text",
                "text": (f"✓ {os.path.basename(fpath)} → {md_filename}\n"
                         f"  图片: {len(image_data_map)} 张提取, {skipped} 张跳过(EMF/WMF/过小)")
            })
        except Exception as e:
            content_parts.append({"type": "text", "text": f"✗ Error parsing {fpath}: {e}"})

    testcase_store["pending_images"] = all_pending_images
    testcase_store["md_files"] = all_md_files

    # Persist to cache
    _save_phase_state("parse", "completed", {
        "file_count": len(all_md_files),
        "total_images": len(all_pending_images)
    })
    _sync_store_to_cache()

    total_imgs = len(all_pending_images)
    summary = (f"\n转换完成:\n"
               f"  Markdown 文件: {len(all_md_files)} 个 → {TMP_DOC_DIR}\n"
               f"  图片文件: {total_imgs} 张 → {TMP_PIC_DIR}\n"
               f"  缓存目录: {TMP_CACHE_DIR}\n")
    if total_imgs > 0:
        summary += (f"\n请逐一调用 get_pending_image 获取待处理图片，\n"
                    f"用视觉能力分析后调用 submit_image_result 提交结果。\n"
                    f"全部处理完成后调用 get_doc_summary 获取文档结构概览。")
    else:
        summary += "\n无需处理图片，可直接调用 get_doc_summary 获取文档结构概览。"

    content_parts.append({"type": "text", "text": summary})
    return {"content": content_parts, "file_count": len(all_md_files), "total_images": total_imgs}

IMAGE_ANALYSIS_PROMPT = (
    "你是一位资深测试开发专家，正在从需求文档中提取测试用例设计所需的信息。\n"
    "请先判断这张图片属于以下哪种类型，然后按对应规则提取具体内容：\n\n"
    "1. 数据表/字段定义 → 用 markdown 表格逐行提取每个字段的：字段名、数据类型、长度、是否必填、默认值、描述。不要遗漏任何一行。\n"
    "2. 流程图/状态机 → 列出所有节点和转换条件，用 A --[条件]--> B 格式描述每条路径（含异常分支），确保无遗漏。\n"
    "3. ER图/架构图 → 列出所有实体及属性，标注实体间关系（一对多等）和外键。\n"
    "4. UI界面/原型图 → 列出所有表单字段（标签、输入类型、可选值）、按钮、表格列头及示例数据。\n"
    "5. 接口/参数定义 → 用 markdown 表格逐个提取参数名、类型、是否必填、取值范围、描述。\n"
    "6. 其他 → 提取所有可见文字和关键信息。\n\n"
    "【输出格式】先用一行标注图片类型，然后输出提取的具体内容。\n"
    "【核心原则】只提取具体数据，禁止笼统概括。看到表格就逐行抄录，看到流程图就逐条列出路径。\n"
    "【测试视角】重点关注：字段约束（长度、格式、必填）、状态转换条件、边界值、业务规则，这些是设计测试用例的关键依据。"
)


def handle_get_pending_image(args):
    pending = testcase_store.get("pending_images", [])
    if not pending:
        # Try restore from cache
        _restore_store_from_cache()
        pending = testcase_store.get("pending_images", [])
    if not pending:
        return {"content": [{"type": "text", "text": "No documents parsed yet. Call parse_documents first."}]}

    next_img = None
    for img in pending:
        if not img["processed"]:
            next_img = img
            break

    if next_img is None:
        _save_phase_state("image_analysis", "completed")
        return {
            "content": [{"type": "text", "text": "所有图片已处理完毕！请调用 get_doc_summary 获取文档结构概览，然后按模块调用 get_doc_section 分段读取内容生成测试用例。"}],
            "all_processed": True
        }

    try:
        with open(next_img["path"], 'rb') as f:
            img_data = f.read()
        b64 = base64.b64encode(img_data).decode('ascii')
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Error reading image {next_img['path']}: {e}"}]}

    total = len(pending)
    processed = sum(1 for p in pending if p["processed"])
    remaining = total - processed - 1

    text_info = (
        f"[{processed + 1}/{total}] 图片ID: {next_img['id']}\n\n"
        f"{IMAGE_ANALYSIS_PROMPT}\n\n"
        f"分析完成后调用 submit_image_result(image_id=\"{next_img['id']}\", analysis=\"你的分析结果\")"
    )

    content_parts = [
        {"type": "text", "text": text_info},
        {"type": "image", "data": b64, "mimeType": next_img["mime"]}
    ]

    result = {
        "content": content_parts,
        "image_id": next_img["id"],
        "total_images": total,
        "processed_count": processed,
        "remaining": remaining + 1,
    }

    # Always show remaining count; no forced session switch.
    # Cross-session resume is still supported — if the system triggers a new
    # session, calling get_workflow_state will pick up right where it left off.
    if remaining > 0:
        content_parts.append({"type": "text", "text": f"提交后还剩 {remaining} 张待处理。"})

    return result

def handle_submit_image_result(args):
    image_id = args.get("image_id", "")
    analysis = args.get("analysis", "")

    if not image_id:
        return {"content": [{"type": "text", "text": "Missing required parameter: image_id"}]}
    if not analysis:
        return {"content": [{"type": "text", "text": "Missing required parameter: analysis"}]}

    pending = testcase_store.get("pending_images", [])
    if not pending:
        _restore_store_from_cache()
        pending = testcase_store.get("pending_images", [])

    target = None
    for img in pending:
        if img["id"] == image_id:
            target = img
            break

    if target is None:
        return {"content": [{"type": "text", "text": f"Image ID not found: {image_id}"}]}
    if target["processed"]:
        return {"content": [{"type": "text", "text": f"Image {image_id} already processed."}]}

    source_doc = target["source_doc"]
    md_path = os.path.join(TMP_DOC_DIR, f"{source_doc}.md")

    if not os.path.exists(md_path):
        return {"content": [{"type": "text", "text": f"Markdown file not found: {md_path}"}]}

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        placeholder = f"{{{{IMG:{image_id}}}}}"
        replacement = f"<!-- 图片分析: {target['filename']} -->\n{analysis}\n<!-- /图片分析 -->"

        if placeholder in md_content:
            md_content = md_content.replace(placeholder, replacement)
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            target["processed"] = True
        else:
            base_id = os.path.splitext(image_id)[0]
            found = False
            for pattern_try in [f"{{{{IMG:{base_id}.png}}}}", f"{{{{IMG:{base_id}.jpg}}}}", f"{{{{IMG:{base_id}.jpeg}}}}"]:
                if pattern_try in md_content:
                    md_content = md_content.replace(pattern_try, replacement)
                    with open(md_path, 'w', encoding='utf-8') as f:
                        f.write(md_content)
                    target["processed"] = True
                    found = True
                    break
            if not found:
                target["processed"] = True
                return {"content": [{"type": "text", "text": f"Warning: placeholder {{{{IMG:{image_id}}}}} not found in {md_path}. Marked as processed."}]}
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Error updating markdown: {e}"}]}

    # Persist progress
    _sync_store_to_cache()
    testcase_store["session_image_count"] += 1
    _save_phase_state("image_analysis", "in_progress", {
        "total": len(pending),
        "processed": sum(1 for p in pending if p["processed"]),
    })

    total = len(pending)
    processed = sum(1 for p in pending if p["processed"])
    remaining = total - processed

    msg = f"✓ 已将图片 [{image_id}] 的分析结果写入 {source_doc}.md ({processed}/{total} 已处理)"
    if remaining > 0:
        msg += f"\n请继续调用 get_pending_image 获取下一张图片。"
    else:
        msg += "\n所有图片已处理完毕！请调用 get_doc_summary 获取文档结构概览。"

    return {
        "content": [{"type": "text", "text": msg}],
        "processed_count": processed,
        "total_images": total,
        "remaining": remaining
    }

def handle_get_workflow_state(args):
    """Return current workflow state for session resume."""
    directory = args.get("directory", os.getcwd())
    _update_workspace(directory)

    # Try to restore from cache
    _restore_store_from_cache()

    state = _load_cache("phase_state.json")
    if not state:
        return {"content": [{"type": "text", "text": "没有找到已保存的工作流状态。请从 parse_documents 开始新的工作流。"}],
                "has_state": False}

    pending = testcase_store.get("pending_images", [])
    total_imgs = len(pending)
    processed_imgs = sum(1 for p in pending if p["processed"])
    unprocessed_imgs = total_imgs - processed_imgs

    md_files = testcase_store.get("md_files", [])
    modules = testcase_store.get("modules", [])
    total_cases = sum(len(s.get("test_cases", [])) for m in modules for s in m.get("sub_modules", []))

    # Determine current phase and next action
    phases = state.get("phases", {})
    lines = ["📋 工作流状态恢复:", ""]

    # Parse phase
    parse_status = phases.get("parse", {}).get("status", "pending")
    lines.append(f"  阶段1 文档解析: {parse_status}")
    if md_files:
        lines.append(f"    - {len(md_files)} 个 Markdown 文件")

    # Image analysis phase
    img_status = phases.get("image_analysis", {}).get("status", "pending")
    lines.append(f"  阶段2 图片分析: {img_status}")
    if total_imgs > 0:
        lines.append(f"    - {processed_imgs}/{total_imgs} 张已处理, {unprocessed_imgs} 张待处理")

    # Generation phase
    gen_status = phases.get("generation", {}).get("status", "pending")
    lines.append(f"  阶段3 用例生成: {gen_status}")
    if modules:
        lines.append(f"    - {len(modules)} 个模块, {total_cases} 个用例")

    # Determine resume instruction
    lines.append("")
    img_completed = img_status == "completed"
    has_unprocessed_images = unprocessed_imgs > 0
    has_testcases = total_cases > 0

    if has_unprocessed_images:
        lines.append(f"▶ 继续操作: 调用 get_pending_image 处理剩余 {unprocessed_imgs} 张图片")
    elif not img_completed and total_imgs > 0:
        # Images exist but status not marked completed yet
        lines.append("▶ 继续操作: 调用 get_pending_image 检查图片处理状态")
    elif (img_completed or total_imgs == 0) and not has_testcases:
        lines.append("▶ 继续操作: 调用 get_doc_summary 获取文档结构，然后按模块生成测试用例")
    elif has_testcases:
        lines.append("▶ 继续操作: 调用 export_xmind 导出最终文件")
    elif parse_status == "completed":
        lines.append("▶ 继续操作: 调用 get_pending_image 开始处理图片")
    else:
        lines.append("▶ 继续操作: 调用 parse_documents 开始新的工作流")

    return {
        "content": [{"type": "text", "text": "\n".join(lines)}],
        "has_state": True,
        "current_phase": state.get("current_phase"),
        "unprocessed_images": unprocessed_imgs,
        "total_cases": total_cases,
        "module_count": len(modules),
    }

def handle_get_doc_summary(args):
    """Return document structure summary without full content."""
    if not testcase_store.get("md_files"):
        _restore_store_from_cache()
    if not testcase_store.get("md_files"):
        return {"content": [{"type": "text", "text": "No markdown files found. Call parse_documents first."}]}

    summary = _build_doc_summary()
    lines = ["📄 文档结构概览:", ""]

    for doc in summary["documents"]:
        lines.append(f"📁 {doc['name']} ({doc['total_chars']} 字符)")
        for sec in doc["sections"]:
            indent = "  " * sec["level"]
            lines.append(f"{indent}{'#' * sec['level']} {sec['heading']} ({sec['char_count']} 字符)")
        lines.append("")

    lines.append(f"总计: {len(summary['documents'])} 个文档, {summary['total_sections']} 个章节, {summary['total_chars']} 字符")
    lines.append("")
    lines.append("请按模块调用 get_doc_section(doc_name, section_heading) 分段读取内容，")
    lines.append("每读取一个模块就生成该模块的测试用例，避免一次性加载全部文档。")

    return {
        "content": [{"type": "text", "text": "\n".join(lines)}],
        "summary": summary,
    }


def handle_get_doc_section(args):
    """Read a specific section from a markdown document."""
    doc_name = args.get("doc_name", "")
    section_heading = args.get("section_heading", "")
    include_sub = args.get("include_subsections", True)

    if not doc_name:
        return {"content": [{"type": "text", "text": "Missing required parameter: doc_name"}]}

    if not testcase_store.get("md_files"):
        _restore_store_from_cache()

    # Find the markdown file
    md_path = None
    for md_info in testcase_store.get("md_files", []):
        if md_info["name"] == doc_name:
            md_path = md_info["path"]
            break

    if not md_path:
        # Try fuzzy match
        for md_info in testcase_store.get("md_files", []):
            if doc_name in md_info["name"] or md_info["name"] in doc_name:
                md_path = md_info["path"]
                doc_name = md_info["name"]
                break

    if not md_path or not os.path.exists(md_path):
        available = [m["name"] for m in testcase_store.get("md_files", [])]
        return {"content": [{"type": "text", "text": f"Document not found: {doc_name}\nAvailable: {', '.join(available)}"}]}

    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Error reading {md_path}: {e}"}]}

    # If no section specified, return full content
    if not section_heading:
        return {"content": [{"type": "text", "text": f"FILE: {doc_name}\n\n{content}"}]}

    lines = content.split('\n')
    sections = _parse_md_sections(content)

    # Find matching section (fuzzy)
    target_sec = None
    for sec in sections:
        if section_heading in sec["heading"] or sec["heading"] in section_heading:
            target_sec = sec
            break
    # Looser match
    if not target_sec:
        section_lower = section_heading.lower()
        for sec in sections:
            if section_lower in sec["heading"].lower() or sec["heading"].lower() in section_lower:
                target_sec = sec
                break

    if not target_sec:
        available_headings = [s["heading"] for s in sections]
        return {"content": [{"type": "text", "text": (
            f"Section not found: '{section_heading}' in {doc_name}\n"
            f"Available sections:\n" + "\n".join(f"  - {h}" for h in available_headings)
        )}]}

    # Determine end line
    start = target_sec["start"]
    if include_sub:
        # Include all child sections (find next section at same or higher level)
        end = len(lines)
        for sec in sections:
            if sec["start"] > start and sec["level"] <= target_sec["level"]:
                end = sec["start"]
                break
    else:
        # Only this section's own content (up to next heading of any level)
        end = target_sec["end"]

    section_content = '\n'.join(lines[start:end])

    return {
        "content": [{"type": "text", "text": f"SECTION: {target_sec['heading']} (from {doc_name})\n\n{section_content}"}],
        "section_heading": target_sec["heading"],
        "char_count": len(section_content),
    }

def handle_get_parsed_markdown(args):
    """Read all markdown files - full content. Prefer get_doc_summary + get_doc_section for large docs."""
    if not testcase_store.get("md_files"):
        _restore_store_from_cache()

    md_files = testcase_store.get("md_files", [])
    if not md_files:
        return {"content": [{"type": "text", "text": "No markdown files found. Call parse_documents first."}]}

    pending = testcase_store.get("pending_images", [])
    unprocessed = sum(1 for p in pending if not p["processed"])

    content_parts = []
    total_chars = 0

    for md_info in md_files:
        try:
            with open(md_info["path"], 'r', encoding='utf-8') as f:
                md_content = f.read()
            total_chars += len(md_content)
            content_parts.append({
                "type": "text",
                "text": f"\n{'='*60}\nFILE: {md_info['name']}\n{'='*60}\n\n{md_content}"
            })
        except Exception as e:
            content_parts.append({"type": "text", "text": f"Error reading {md_info['path']}: {e}"})

    summary = f"\n共 {len(md_files)} 个文档已加载 ({total_chars} 字符)。"
    if unprocessed > 0:
        summary += f"\n注意: 还有 {unprocessed} 张图片未处理。"
    if total_chars > 30000:
        summary += "\n⚠️ 文档内容较大，建议使用 get_doc_summary + get_doc_section 分段读取以减少上下文占用。"

    content_parts.append({"type": "text", "text": summary})
    return {"content": content_parts, "file_count": len(md_files), "total_chars": total_chars}


def handle_save_testcases(args):
    modules = args.get("modules", None)
    append_module = args.get("append_module", None)

    if append_module:
        # Incremental: append one module (replace if same name exists)
        if not testcase_store["modules"]:
            _restore_store_from_cache()
        # Replace existing module with same name, or append new
        mod_name = append_module.get("name", "")
        replaced = False
        for i, existing in enumerate(testcase_store["modules"]):
            if existing.get("name") == mod_name:
                testcase_store["modules"][i] = append_module
                replaced = True
                break
        if not replaced:
            testcase_store["modules"].append(append_module)
        _sync_store_to_cache()
        _save_phase_state("generation", "in_progress", {
            "module_count": len(testcase_store["modules"])
        })
        total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"] for s in m.get("sub_modules", []))
        action = "替换" if replaced else "追加"
        return {"content": [{"type": "text", "text": (
            f"✓ {action}模块 '{mod_name}', "
            f"当前共 {len(testcase_store['modules'])} 个模块, {total} 个用例"
        )}]}

    if modules is None:
        return {"content": [{"type": "text", "text": "Missing parameter: modules or append_module"}]}

    testcase_store["modules"] = modules
    _sync_store_to_cache()
    _save_phase_state("generation", "completed", {"module_count": len(modules)})

    total = sum(len(s.get("test_cases", [])) for m in modules for s in m.get("sub_modules", []))
    return {"content": [{"type": "text", "text": f"Saved {len(modules)} modules, {total} test cases. 已持久化到 .tmp/cache/testcases.json"}]}


def handle_get_testcases(args):
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"]
                for s in m.get("sub_modules", []))
    return {
        "content": [{"type": "text", "text": json.dumps(testcase_store["modules"], ensure_ascii=False, indent=2)}],
        "module_count": len(testcase_store["modules"]),
        "total_cases": total,
    }


def handle_export_xmind(args):
    if not testcase_store["modules"]:
        _restore_store_from_cache()
    if not testcase_store["modules"]:
        return {"content": [{"type": "text", "text": "No test cases to export."}]}
    p = args.get("output_path", os.path.join(_workspace(), "test_cases.xmind"))
    try:
        create_xmind_file(testcase_store["modules"], p)
        total = sum(len(s.get("test_cases", [])) for m in testcase_store["modules"]
                    for s in m.get("sub_modules", []))
        _save_phase_state("export", "completed")
        return {"content": [{"type": "text", "text": f"Exported: {p}\n{len(testcase_store['modules'])} modules, {total} cases"}]}
    except Exception as e:
        return {"content": [{"type": "text", "text": f"Export failed: {e}"}]}

# ============================================================
# MCP Main Loop
# ============================================================

HANDLERS = {
    "setup_environment": handle_setup_environment,
    "parse_documents": handle_parse_documents,
    "get_pending_image": handle_get_pending_image,
    "submit_image_result": handle_submit_image_result,
    "get_workflow_state": handle_get_workflow_state,
    "get_doc_summary": handle_get_doc_summary,
    "get_doc_section": handle_get_doc_section,
    "get_parsed_markdown": handle_get_parsed_markdown,
    "save_testcases": handle_save_testcases,
    "get_testcases": handle_get_testcases,
    "export_xmind": handle_export_xmind,
}


def handle_request(req):
    method = req.get("method", "")
    rid = req.get("id")
    params = req.get("params", {})

    sys.stderr.write(f"[MCP] Received: method={method} id={rid}\n")
    sys.stderr.flush()

    if method == "initialize":
        send_response(rid, {
            "protocolVersion": "2024-11-05",
            "capabilities": {"tools": {"listChanged": False}},
            "serverInfo": {"name": "testcase-generator", "version": "6.0.0"}
        })
    elif method == "notifications/initialized":
        pass
    elif method == "tools/list":
        send_response(rid, {"tools": TOOLS})
    elif method == "tools/call":
        name = params.get("name", "")
        handler = HANDLERS.get(name)
        if handler:
            try:
                result = handler(params.get("arguments", {}))
                send_response(rid, result)
            except Exception as e:
                send_response(rid, {"content": [{"type": "text", "text": f"Error: {e}"}], "isError": True})
        else:
            send_error(rid, -32601, f"Unknown tool: {name}")
    elif method == "ping":
        send_response(rid, {})
    else:
        if rid is not None:
            send_error(rid, -32601, f"Method not found: {method}")


def main():
    sys.stderr.write("TestCase Generator MCP Server v6.0 starting...\n")
    sys.stderr.flush()

    while True:
        try:
            line = sys.stdin.buffer.readline()
            if not line:
                return

            stripped = line.strip()
            if not stripped:
                continue

            decoded = stripped.decode('utf-8', errors='replace')

            if decoded.startswith('{'):
                try:
                    request = json.loads(decoded)
                    handle_request(request)
                except json.JSONDecodeError as e:
                    sys.stderr.write(f"JSON parse error: {e}\n")
                    sys.stderr.flush()
            elif decoded.lower().startswith('content-length'):
                try:
                    cl = int(decoded.split(":", 1)[1].strip())
                except (ValueError, IndexError):
                    continue
                sys.stdin.buffer.readline()  # skip blank line
                body = b""
                while len(body) < cl:
                    chunk = sys.stdin.buffer.read(cl - len(body))
                    if not chunk:
                        break
                    body += chunk
                if body:
                    try:
                        request = json.loads(body.decode('utf-8'))
                        handle_request(request)
                    except json.JSONDecodeError as e:
                        sys.stderr.write(f"JSON parse error: {e}\n")
                        sys.stderr.flush()
        except Exception as e:
            sys.stderr.write(f"Error in main loop: {e}\n")
            sys.stderr.flush()


if __name__ == "__main__":
    main()
